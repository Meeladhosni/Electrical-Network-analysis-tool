import tkinter as tk
from tkinter import ttk, simpledialog, filedialog
import tkinter.font as tkf
from PIL import Image, ImageTk
import numpy as np
from tooltip import ToolTip
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
import datetime
import glob
import os
import subprocess
import threading

import Data_Store
from Calculations import Y_bus_class
from Calculations import Calc_class
import Power_flow
import PLF_calc
import Plots


class ImageResizer:
    @staticmethod
    def resize_image(canvas, img, event):
        # Delete the previous canvas image
        canvas.delete("bg_img")
        canvas.delete("text_tag")
        # Resize the image with LANCZOS filter
        new_width = event.width
        new_height = event.height
        resized_image = img.resize((new_width, new_height))
        photo = ImageTk.PhotoImage(resized_image)

        # Put the image on the canvas and set it as the new background
        canvas.background = photo  # Keep a reference!
        canvas.create_image(0, 0, image=photo, anchor='nw', tags="bg_img")

        # Keep a reference to the photo to prevent garbage collection
        canvas.photo = photo


class GUI_setup:
    def __init__(self, app):

        (self.scale_factor, self.notebook, self.style,
         self.intro_tab, self.data_entry_tab, self.Result_tab) = None, None, None, None, None, None
        self.app = app
        self.background_image_path = os.path.join('config', 'background.png')

        self.original_image = Image.open(self.background_image_path)
        Data_Store.central_data_store.update_data('app', self.app)
        Data_Store.central_data_store.update_data('original_image', self.original_image)

        self.base_width, self.base_height = 1463, 823
        self.base_font_sizes = [10, 12, 14, 18, 20, 30]  # Base font sizes
        self.fonts = [tkf.Font(size=size) for size in self.base_font_sizes]

    def setup_styles(self):
        self.app.state('zoomed')
        self.app.resizable(False, False)
        self.app.title("Electrical Network Analysis Tool")

        icon_path = os.path.join('config', 'icon.ico')
        icon = tk.PhotoImage(icon_path)

        self.app.iconbitmap(True, icon)

        # Set the initial font sizes based on the window size
        screen_width = self.app.winfo_screenwidth()
        screen_height = self.app.winfo_screenheight()
        width = int(screen_width * 0.9)
        height = int(screen_height * 0.9)
        self.app.geometry(f"{width}x{height}")
        self.style = ttk.Style()
        self.style.theme_use('xpnative')

        ratio_x = width / self.base_width
        ratio_y = height / self.base_height

        if ratio_x == 1 and ratio_y == 1:
            scale_x, scale_y = 1, 1
        elif ratio_x > 1 and ratio_y > 1:
            scale_x = np.log2(np.abs(ratio_x))
            scale_x = ratio_x - scale_x / 1.618

            scale_y = np.log2(np.abs(ratio_y))
            scale_y = ratio_y - scale_y / 1.618

        else:
            scale_x, scale_y = ratio_x, ratio_y

        self.scale_factor = min(scale_x, scale_y)

        for i, base_size in enumerate(self.base_font_sizes):
            new_size = max(1, int(base_size * self.scale_factor))
            # Minimum font size of 1
            self.fonts[i].configure(size=new_size)

        Data_Store.central_data_store.update_data('scale_factor', self.scale_factor)
        Data_Store.central_data_store.update_data('fonts', self.fonts)

        # Create and grid the notebook (tabbed interface)
        self.notebook = ttk.Notebook(self.app)
        self.app.grid_rowconfigure(0, weight=1)
        self.app.grid_columnconfigure(0, weight=1)
        self.notebook.grid(row=0, column=0, sticky='nsew', padx=10, pady=10)
        self.intro_tab = ttk.Frame(self.notebook)
        self.data_entry_tab = ttk.Frame(self.notebook)
        self.Result_tab = ttk.Frame(self.notebook)

        self.notebook.add(self.intro_tab, text="data 1")

        intro_canvas = tk.Canvas(self.intro_tab)
        intro_canvas.place(relwidth=1, relheight=1)
        intro_canvas.bind('<Configure>',
                          lambda event, c=intro_canvas, img=self.original_image:
                          ImageResizer.resize_image(c, img, event))
        intro_canvas.original_image = self.original_image

        data_entry_canvas = tk.Canvas(self.data_entry_tab)
        data_entry_canvas.place(relwidth=1, relheight=1)
        data_entry_canvas.bind('<Configure>',
                               lambda event, c=data_entry_canvas, img=self.original_image:
                               ImageResizer.resize_image(c, img, event))
        data_entry_canvas.original_image = self.original_image

        Result_canvas = tk.Canvas(self.Result_tab)
        Result_canvas.place(relwidth=1, relheight=1)
        Result_canvas.bind('<Configure>',
                           lambda event, c=Result_canvas, img=self.original_image:
                           ImageResizer.resize_image(c, img, event))
        Result_canvas.original_image = self.original_image

        # Add a scrollbar to the frame
        scrollbar_ttk = ttk.Scrollbar(self.data_entry_tab, orient='vertical')
        scrollbar_ttk.pack(side='right', fill='y')

        # Create a canvas in the frame and attach the scrollbar to it
        scrollbar_canvas = tk.Canvas(self.data_entry_tab, yscrollcommand=scrollbar_ttk.set)
        scrollbar_canvas.pack(side='right', fill='both', expand=True)

        # Configure the scrollbar
        scrollbar_ttk.config(command=scrollbar_canvas.yview)

        # Create a frame inside the canvas to hold the contents
        scrollbar_frame = ttk.Frame(scrollbar_canvas)
        scrollbar_canvas.create_window((0, 0), window=scrollbar_frame, anchor='nw')

        # Update the scroll region when the size of the frame changes
        scrollbar_frame.bind('<Configure>',
                             lambda e: scrollbar_canvas.configure(scrollregion=scrollbar_canvas.bbox('all')))

        # Add some content to the content_frame for demonstration
        for length in range(50):
            ttk.Label(scrollbar_frame, text=f'').pack()

    def add_data_entry_tab(self):
        self.notebook.add(self.data_entry_tab, text="data_entry_tab")
        for i in range(self.notebook.index('end')):
            self.notebook.hide(i)
        self.notebook.select(self.data_entry_tab)

    def add_result_tab(self):
        self.notebook.add(self.Result_tab, text="Result_tab")
        for i in range(self.notebook.index('end')):
            self.notebook.hide(i)
        self.notebook.select(self.Result_tab)

    def hide_all_tabs(self, previous_canvas):
        for i in range(self.notebook.index('end')):
            self.notebook.hide(i)
        self.notebook.select(previous_canvas)

    def hide_all_tabs_1(self):
        # Get the total number of tabs
        total_tabs = self.notebook.index('end')

        # Iterate through each tab
        for i in range(total_tabs):
            # Check if the current tab is not the one to keep
            if i != self.intro_tab:
                self.notebook.hide(i)

        self.notebook.select(self.intro_tab)

    def open_all_tabs(self):
        self.notebook.add(self.data_entry_tab, text="data_entry_tab")
        self.notebook.add(self.Result_tab, text="Result_tab")


class intro_class:
    def __init__(self, canvas):
        self.Placement_hard_reset_btn, self.open_all_btn, self.num_buses, self.method = None, None, None, None
        self.master = canvas
        self.canvas = tk.Canvas(self.master)
        self.canvas.place(relwidth=1, relheight=1)
        self.app = Data_Store.central_data_store.get_data('app')
        self.gui_setup = Data_Store.central_data_store.get_data('gui_setup')
        self.original_image = Data_Store.central_data_store.get_data('original_image')
        self.fonts = Data_Store.central_data_store.get_data('fonts')
        self.scale_factor = Data_Store.central_data_store.get_data('scale_factor')
        self.saved_x_value = Data_Store.central_data_store.get_data('saved_x_value')
        self.saved_y_value = Data_Store.central_data_store.get_data('saved_y_value')
        self.saved_f_value = Data_Store.central_data_store.get_data('saved_f_value')
        self.ref_x = self.saved_x_value * self.scale_factor
        self.ref_y = self.saved_y_value * self.scale_factor
        self.ref_f = self.saved_f_value
        self.frame_x_Placement, self.frame_x_Placement_copy, self.Pos_btn, self.ref_data = None, None, None, None
        self.place_font_ref_file_path = os.path.join('Config', 'place_font_ref_info.txt')
        self.dyn_font_var, self.Placement_frame, self.Placement_reset_btn = None, None, None
        self.x_scale, self.y_scale, self.f_scale = None, None, None
        self.x_scale_value, self.y_scale_value, self.f_scale_value = (
            self.saved_x_value, self.saved_y_value, self.saved_f_value)

        (self.Placement_x_var, self.Placement_y_var, self.Placement_show_btn) = None, None, None
        self.methods_list = ['- GS', '- NR', '- NRFD', '- PLF-NR', '- PLF-NRFD']
        self.num_buses_space = tk.IntVar()
        self.methods_listbox, self.selected_signature_initials, self.Placement_frame_window = None, None, None
        self.f12, self.f14, self.f18, self.f30 = self.fonts[1], self.fonts[2], self.fonts[3], self.fonts[5]

        self.canvas.bind('<Configure>', self.intro_setup)

    def intro_setup(self, event):
        # Delete the previous canvas image
        self.canvas.delete("bg_img")
        new_width = event.width
        new_height = event.height
        resized_image = self.original_image.resize((new_width, new_height))
        photo = ImageTk.PhotoImage(resized_image)

        self.canvas.background = photo  # reference!
        self.canvas.create_image(0, 0, image=photo, anchor='nw', tags="bg_img")

        self.frame_x_Placement = 0.998 * new_width  # Initial position
        self.frame_x_Placement_copy = 0.998 * new_width  # Initial position
        self.setup_Placement_frame()
        self.intro_widgets()

    def intro_widgets(self):
        self.canvas.delete("text_tag")
        self.canvas.delete("error_message")
        self.canvas.create_text(8 * self.ref_x,
                                1 * self.ref_y,
                                text="Hi there! I'm Power Flow Analysis Tool",
                                tags="text_tag", font=('Verdana', int(self.f30.cget('size') + self.ref_f), 'bold'),
                                justify="center")

        # number of buses:
        self.canvas.create_text(4.9 * self.ref_x,
                                2.2 * self.ref_y,
                                text="Number of buses:", tags="text_tag",
                                font=('Verdana', int(self.f18.cget('size') + self.ref_f), 'bold'), justify="left")

        num_buses_widget = ttk.Entry(self.canvas, textvariable=self.num_buses_space,
                                     width=3, justify="center",
                                     font=('Verdana', int(self.f18.cget('size') + self.ref_f)))
        self.canvas.create_window(8 * self.ref_x,
                                  2.35 * self.ref_y,
                                  window=num_buses_widget, tags="text_tag")

        # methods:
        self.canvas.create_text(4.90 * self.ref_x,
                                3.24 * self.ref_y,
                                text="Available methods:", tags="text_tag",
                                font=('Verdana', int(self.f18.cget('size') + self.ref_f), 'bold'), justify="left")

        self.canvas.create_text(14 * self.ref_x,
                                3.6 * self.ref_y,
                                text="- (GS) Gauss Siedel\n- (NR) Newton Raphson\n"
                                     "- (NRFD) Newton Raphson Fast Decoupled\n"
                                     "- (PLF-NR) Probabilistic Load Flow with NR\n"
                                     "- (PLF-NRFD) Probabilistic Load Flow with NRFD", tags="text_tag",
                                font=('Verdana', int(self.f12.cget('size') + self.ref_f), 'bold'), justify="left")

        self.methods_listbox = tk.Listbox(self.canvas, font=('Verdana', int(self.f14.cget('size') + self.ref_f)),
                                          height=5, width=16, exportselection=False)
        for item in self.methods_list:
            self.methods_listbox.insert(tk.END, item)
        methods_scrollbar = tk.Scrollbar(self.canvas,
                                         orient="vertical", command=self.methods_listbox.yview)
        self.methods_listbox.configure(yscrollcommand=methods_scrollbar.set)
        self.canvas.create_window(7.5 * self.ref_x,
                                  3.6 * self.ref_y,
                                  window=self.methods_listbox, tags="text_tag")

        self.canvas.create_window(8.3 * self.ref_x,
                                  3.6 * self.ref_y,
                                  window=methods_scrollbar, height=95, tags="text_tag")

        style = ttk.Style()
        style.configure("Custom.TButton", font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
        Next_btn = ttk.Button(self.canvas, text="Next",
                              command=lambda: self.intro_get(),
                              style="Custom.TButton")
        self.canvas.create_window(9 * self.ref_x,
                                  5.2 * self.ref_y,
                                  window=Next_btn, tags="text_tag", )
        ToolTip(Next_btn, text="Proceed with selected method & num of buses", justify="center")


        Load_btn = ttk.Button(self.canvas, text="📂 Load Data",
                              command=lambda: self.load_data(), style="Custom.TButton")
        self.canvas.create_window(6 * self.ref_x, 5.2 * self.ref_y, window=Load_btn, tags="text_tag")
        ToolTip(Load_btn, text="Proceed with selected method & saved data", justify="center")

    def check_input_error(self, attribute_name):
        try:
            float(getattr(self, attribute_name).get())
            return True
        except ValueError:
            # Catch ValueError if float conversion fails due to incorrect string formats (non-numeric)
            return False
        except tk.TclError:
            # Catch TclError if tkinter variable handling fails (e.g., bad string passed to DoubleVar)
            return False

    def intro_get(self):
        loaded_data_switch = 0
        num_buses_input = self.check_input_error(f"num_buses_space")
        if num_buses_input:
            self.num_buses = int(self.num_buses_space.get())
        else:
            self.display_error_message("Please enter a valid number")
            return
        method_selected_index = self.methods_listbox.curselection()
        if method_selected_index:
            self.method = self.methods_list[method_selected_index[0]]

        if not self.num_buses or not self.method:
            self.display_error_message("Please fill in all required fields.")
            return
        else:
            self.display_error_message("")
        Data_Store.central_data_store.update_data('num_buses', self.num_buses)
        Data_Store.central_data_store.update_data('method', self.method)
        Data_Store.central_data_store.update_data('loaded_data_switch', loaded_data_switch)
        Data_Store.central_data_store.update_data('ref_x', self.ref_x)
        Data_Store.central_data_store.update_data('ref_y', self.ref_y)
        Data_Store.central_data_store.update_data('ref_f', self.ref_f)

        self.gui_setup.add_data_entry_tab()
        self.save_scale()

        # Run next class:
        Entry_class(self.gui_setup.data_entry_tab, self.gui_setup.intro_tab)

    def load_data(self):
        method_selected_index = self.methods_listbox.curselection()
        if method_selected_index:
            self.method = self.methods_list[method_selected_index[0]]

        if not self.method:
            self.display_error_message("Please choose a method first.")
            return
        else:
            self.display_error_message("")

        loaded_data_switch = 1
        processor = Data_Store.DataLdr(self.app)
        processor.choose_file()

        loaded_var = processor.read_file()
        if loaded_var == "Please choose a file first.":
            self.display_error_message("Please choose a file first.")
            return
        else:
            self.display_error_message("")

        num_buses = loaded_var.get("num_buses")
        Data_Store.central_data_store.update_data('num_buses', num_buses)
        Data_Store.central_data_store.update_data('R_values', loaded_var.get("R_values"))
        Data_Store.central_data_store.update_data('X_values', loaded_var.get("X_values"))
        Data_Store.central_data_store.update_data('B_values', loaded_var.get("B_values"))
        Data_Store.central_data_store.update_data('T_values', loaded_var.get("T_values"))
        Data_Store.central_data_store.update_data('Gs_array', loaded_var.get("Gs_array"))
        Data_Store.central_data_store.update_data('Bs_array', loaded_var.get("Bs_array"))
        Data_Store.central_data_store.update_data('V_array', loaded_var.get("V_array"))
        Data_Store.central_data_store.update_data('PG_array', loaded_var.get("PG_array"))
        Data_Store.central_data_store.update_data('QG_array', loaded_var.get("QG_array"))
        Data_Store.central_data_store.update_data('PL_array', loaded_var.get("PL_array"))
        Data_Store.central_data_store.update_data('QL_array', loaded_var.get("QL_array"))
        Data_Store.central_data_store.update_data('Q_min_array', loaded_var.get("Q_min_array"))
        Data_Store.central_data_store.update_data('Q_max_array', loaded_var.get("Q_max_array"))
        Data_Store.central_data_store.update_data('bus_type_array', loaded_var.get("bus_type_array"))

        for i in range(1, num_buses + 1):
            PG_var = f"PG_sub_{i}"
            Data_Store.central_data_store.update_data(f"PG_sub_{i}", loaded_var.get(PG_var))

            QG_var = f"QG_sub_{i}"
            Data_Store.central_data_store.update_data(f"QG_sub_{i}", loaded_var.get(QG_var))

        if self.method == '- PLF-NR' or self.method == '- PLF-NRFD':
            std_dev_array = loaded_var.get("std_dev_array")
            Data_Store.central_data_store.update_data('std_dev_array', std_dev_array)

        Data_Store.central_data_store.update_data('method', self.method)
        Data_Store.central_data_store.update_data('loaded_data_switch', loaded_data_switch)
        Data_Store.central_data_store.update_data('ref_x', self.ref_x)
        Data_Store.central_data_store.update_data('ref_y', self.ref_y)
        Data_Store.central_data_store.update_data('ref_f', self.ref_f)
        self.gui_setup.add_data_entry_tab()

        # Run next class:
        Entry_class(self.gui_setup.data_entry_tab, self.gui_setup.intro_tab)
        self.save_scale()

    def display_error_message(self, message):
        self.canvas.delete("error_text")
        Dark_red = '#8b0000'
        self.canvas.create_text(10 * self.ref_x, 6.075 * self.ref_y, text=message, fill=Dark_red, tags="error_text",
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))

    def toggle_Placement_frame(self):
        if self.Placement_show_btn.cget("text") == "<":
            self.slide_Placement_frame(1 * self.ref_y, -7)
            self.Placement_show_btn.config(text=">")
            self.Pos_btn["state"] = "disabled"

        elif self.Placement_show_btn.cget("text") == ">":
            self.slide_Placement_frame(self.frame_x_Placement_copy, 7)
            self.Placement_show_btn.config(text="<")
            self.Pos_btn["state"] = "normal"

    def slide_Placement_frame(self, end_x, step):
        if (step > 0 and self.frame_x_Placement < end_x) or (step < 0 and self.frame_x_Placement > end_x):
            self.frame_x_Placement += step
            self.canvas.move(self.Placement_frame_window, step, 0)
            self.app.after(2, lambda: self.slide_Placement_frame(end_x, step))

    def on_scale_change_x(self, x_scale_value):
        self.x_scale_value = x_scale_value
        self.ref_x = float(x_scale_value) * self.scale_factor
        self.intro_widgets()

    def on_scale_change_y(self, y_scale_value):
        self.y_scale_value = y_scale_value
        self.ref_y = float(y_scale_value) * self.scale_factor
        self.intro_widgets()

    def on_scale_change_font(self, f_scale_value):
        self.f_scale_value = f_scale_value
        self.ref_f = float(f_scale_value)
        self.intro_widgets()

    def reset_scale(self):
        # Reset the scale to last saved value
        self.Placement_x_var.set(self.saved_x_value)
        self.Placement_y_var.set(self.saved_y_value)
        self.dyn_font_var.set(self.saved_f_value)

        self.ref_x = self.saved_x_value * self.scale_factor
        self.ref_y = self.saved_y_value * self.scale_factor
        self.ref_f = self.saved_f_value

        self.intro_widgets()

    def reset_scale_to_def(self):
        self.Placement_x_var.set(100)
        self.Placement_y_var.set(100)
        self.dyn_font_var.set(0)
        self.ref_x, self.ref_y, self.ref_f = 100 * self.scale_factor, 100 * self.scale_factor, 0

        self.ref_data = f"saved_x_value = {100}\nsaved_y_value = {100}\nsaved_f_value = {0}"
        with open(self.place_font_ref_file_path, 'w') as file:
            file.write(self.ref_data)

        self.intro_widgets()

    def save_scale(self):
        self.ref_data = (
            f"saved_x_value = {self.x_scale_value}\n"
            f"saved_y_value = {self.y_scale_value}\n"
            f"saved_f_value = {self.f_scale_value}\n"
        )
        with open(self.place_font_ref_file_path, 'w') as file:
            file.write(self.ref_data)

    def setup_Placement_frame(self):
        self.Placement_frame = tk.Frame(self.app, width=int(round(500 * self.scale_factor)),
                                        height=int(round(200 * self.scale_factor)))

        style = ttk.Style()
        style.configure("Custom.TButton",
                        font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
        self.Pos_btn = ttk.Button(self.canvas, text="⌖",
                                  command=lambda: self.toggle_Placement_frame(),
                                  style="Custom.TButton", width=2)
        self.canvas.create_window(11, 6, window=self.Pos_btn, tags="pos_tag")
        ToolTip(self.Pos_btn, text="Change Position and Font", justify="center")

        self.Placement_show_btn = ttk.Button(self.Placement_frame, text="<",
                                             command=self.toggle_Placement_frame, width=1)
        self.Placement_show_btn.pack(side=tk.LEFT, fill=tk.Y, expand=True)

        self.Placement_x_var = tk.DoubleVar(value=self.saved_x_value)
        self.x_scale = tk.Scale(self.Placement_frame, from_=50, to=150,
                                orient=tk.HORIZONTAL, variable=self.Placement_x_var,
                                command=self.on_scale_change_x)
        self.x_scale.pack(side=tk.TOP)
        ToolTip(self.x_scale, text="X Position", justify="center")

        self.Placement_y_var = tk.DoubleVar(value=self.saved_y_value)
        self.y_scale = tk.Scale(self.Placement_frame, from_=50, to=150,
                                orient=tk.HORIZONTAL, variable=self.Placement_y_var,
                                command=self.on_scale_change_y)
        self.y_scale.pack(side=tk.TOP)
        ToolTip(self.y_scale, text="Y Position", justify="center")

        self.dyn_font_var = tk.DoubleVar(value=self.ref_f)
        self.f_scale = tk.Scale(self.Placement_frame, from_=-10, to=10,
                                orient=tk.HORIZONTAL, variable=self.dyn_font_var,
                                command=self.on_scale_change_font)
        self.f_scale.pack(side=tk.TOP)
        ToolTip(self.f_scale, text="Font Size", justify="center")

        self.Placement_frame_window = self.canvas.create_window(self.frame_x_Placement, self.ref_y * 0.04,
                                                                window=self.Placement_frame, anchor='nw')

        self.open_all_btn = ttk.Button(self.Placement_frame, text="Show all/ Refresh", command=self.open_all_tabs)
        self.open_all_btn.pack(side=tk.BOTTOM, fill=tk.X)

        self.Placement_hard_reset_btn = ttk.Button(self.Placement_frame,
                                                   text="Default", command=self.reset_scale_to_def)
        self.Placement_hard_reset_btn.pack(side=tk.LEFT, fill=tk.X)

        self.Placement_reset_btn = ttk.Button(self.Placement_frame, text="Last saved", command=self.reset_scale)
        self.Placement_reset_btn.pack(side=tk.RIGHT, fill=tk.X)

    def open_all_tabs(self):
        Data_Store.central_data_store.update_data('ref_x', self.ref_x)
        Data_Store.central_data_store.update_data('ref_y', self.ref_y)
        Data_Store.central_data_store.update_data('ref_f', self.ref_f)
        Data_Store.central_data_store.update_data('num_buses', 5)
        Data_Store.central_data_store.update_data('method', "dyn Config")
        Data_Store.central_data_store.update_data('iteration_result', "dynamic Config")
        Data_Store.central_data_store.update_data('bus_type_array', ['Slack', 'PV', 'PQ', 'PQ', 'PQ'])
        Data_Store.central_data_store.update_data('V_mag_output_array', [1, 1, 1, 1, 1])
        Data_Store.central_data_store.update_data('V_ang_output_array', [1, 1, 1, 1, 1])
        Data_Store.central_data_store.update_data('PG_output_array', [1, 1, 1, 1, 1])
        Data_Store.central_data_store.update_data('QG_output_array', [1, 1, 1, 1, 1])
        Data_Store.central_data_store.update_data('PL_array', [1, 1, 1, 1, 1])
        Data_Store.central_data_store.update_data('QL_array', [1, 1, 1, 1, 1])
        Data_Store.central_data_store.update_data('total_P_loss', 1)
        Data_Store.central_data_store.update_data('total_Q_loss', 1)
        Data_Store.central_data_store.update_data('total_P_generation', 1)
        Data_Store.central_data_store.update_data('total_P_load_loss', 1)
        Data_Store.central_data_store.update_data('I_ij_array', [1 + 1j, 1 + 1j, 1 + 1j, 1 + 1j, 1 + 1j])
        Data_Store.central_data_store.update_data('S_ij_array', [1 + 1j, 1 + 1j, 1 + 1j, 1 + 1j, 1 + 1j])
        Data_Store.central_data_store.update_data('S_ij_losses_array', [1 + 1j, 1 + 1j, 1 + 1j, 1 + 1j, 1 + 1j])
        self.gui_setup.open_all_tabs()
        Entry_class(self.gui_setup.data_entry_tab, self.gui_setup.intro_tab)
        Result_class(self.gui_setup.Result_tab, self.gui_setup.intro_tab)


class Entry_class:
    def __init__(self, master, previous_canvas):
        self.image_info, self.tip_windows, self.anchor_point = None, None, None
        self.tip_window, self.image_path, self.size = None, None, None
        self.tolerance, self.max_iterations, self.Processing_text = None, None, None
        self.slack_bus_num, self.gen_num, self.thread_end, self.thread_error_end = None, None, False, False
        self.Back_btn, self.save_btn, self.Run_btn, self.btn_list = None, None, None, None
        self.Processing_after, self.Processing_dots, self.initial_top = None, None, None
        self.max_scroll_height = 0
        self.master = master
        self.previous_canvas = previous_canvas
        self.canvas = tk.Canvas(self.master, bg='#AFAFAF')
        self.canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        self.app = Data_Store.central_data_store.get_data('app')
        self.gui_setup = Data_Store.central_data_store.get_data('gui_setup')
        self.original_image = Data_Store.central_data_store.get_data('original_image')
        self.fonts = Data_Store.central_data_store.get_data('fonts')
        self.ref_x = Data_Store.central_data_store.get_data('ref_x')
        self.ref_y = Data_Store.central_data_store.get_data('ref_y')
        self.ref_f = Data_Store.central_data_store.get_data('ref_f')
        self.scrollbar = tk.Scrollbar(self.master, orient="vertical", command=self.canvas.yview)
        self.scrollbar.place(relx=0.99, rely=0, relwidth=0.01, relheight=1)

        self.loaded_data_switch = Data_Store.central_data_store.get_data('loaded_data_switch')
        self.method = Data_Store.central_data_store.get_data('method')
        self.num_buses = Data_Store.central_data_store.get_data('num_buses')
        self.R_values, self.X_values, self.B_values, self.T_values, self.Z_values = {}, {}, {}, {}, {}
        self.Gs_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.Bs_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.PG_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.QG_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.PL_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.QL_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.Q_min_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.Q_max_array = np.array([0.0] * self.num_buses, dtype=np.float64)
        self.V_array = np.array([complex(1.0, 0.0)] * self.num_buses, dtype=np.complex128)
        self.bus_type_array = np.array(['empty' for _ in range(self.num_buses)])
        self.std_dev_array = np.array([0.1] * 4)

        for i in range(1, self.num_buses):
            for j in range(i + 1, self.num_buses + 1):
                pair = f"{i}_{j}"
                for prefix in ["R", "X", "B", "T"]:
                    var_name = f"{prefix}{pair}_space"
                    setattr(self, var_name, tk.DoubleVar(value=0.0))

        for i in range(self.num_buses):
            Var_name = f"V{i + 1}_space"
            setattr(self, Var_name, tk.DoubleVar(value=1.0))

        properties = ["ang", "PG", "QG", "PL", "QL", "Q_min", "Q_max", "Gs", "Bs"]

        for i in range(self.num_buses):
            for prop in properties:
                Var_name = f"{prop}{i + 1}_space"
                setattr(self, Var_name, tk.DoubleVar(value=0.0))

        self.slack_bus_option = tk.IntVar(value=-1)
        self.bus_options = [tk.IntVar(value=-1) for _ in range(self.num_buses)]
        self.max_iterations_space = tk.IntVar(value=200)
        self.tolerance_space = tk.IntVar(value=8)
        self.slack_widgets, self.PV_widgets, self.PQ_widgets = [], [], []

        self.f10, self.f12, self.f14, self.f20 = self.fonts[0], self.fonts[1], self.fonts[2], self.fonts[4]
        self.setup_radio_buttons()

        std_0 = std_1 = std_2 = std_3 = 0.1
        self.std_dev_PL_space = tk.DoubleVar(value=std_0)
        self.std_dev_QL_space = tk.DoubleVar(value=std_1)
        self.std_dev_PG_space = tk.DoubleVar(value=std_2)
        self.std_dev_QG_space = tk.DoubleVar(value=std_3)
        self.num_simulations_space = tk.IntVar(value=1000)
        self.PG_array_loaded, self.QG_array_loaded = None, None
        if self.loaded_data_switch == 1:
            R_values_loaded = Data_Store.central_data_store.get_data('R_values')
            X_values_loaded = Data_Store.central_data_store.get_data('X_values')
            B_values_loaded = Data_Store.central_data_store.get_data('B_values')
            T_values_loaded = Data_Store.central_data_store.get_data('T_values')
            Gs_array_loaded = Data_Store.central_data_store.get_data('Gs_array')
            Bs_array_loaded = Data_Store.central_data_store.get_data('Bs_array')
            self.PG_array_loaded = Data_Store.central_data_store.get_data('PG_array')
            self.QG_array_loaded = Data_Store.central_data_store.get_data('QG_array')
            PL_array_loaded = Data_Store.central_data_store.get_data('PL_array')
            QL_array_loaded = Data_Store.central_data_store.get_data('QL_array')
            Q_min_array_loaded = Data_Store.central_data_store.get_data('Q_min_array')
            Q_max_array_loaded = Data_Store.central_data_store.get_data('Q_max_array')
            V_array_loaded = Data_Store.central_data_store.get_data('V_array')
            bus_type_array_loaded = Data_Store.central_data_store.get_data('bus_type_array')

            for pair in R_values_loaded.keys():
                # Retrieve the loaded values or use default if key not found
                R_value_default = R_values_loaded.get(pair, 0.0)
                X_value_default = X_values_loaded.get(pair, 0.0)
                B_value_default = B_values_loaded.get(pair, 0.0)
                T_value_default = T_values_loaded.get(pair, 0.0)

                # Overwrite the tk.DoubleVar values
                for prefix, value in zip(["R", "X", "B", "T"],
                                         [R_value_default, X_value_default, B_value_default, T_value_default]):
                    var_name = f"{prefix}{pair}_space"
                    # Overwrite the existing tk.DoubleVar variable with the new value
                    if hasattr(self, var_name):
                        getattr(self, var_name).set(value)

            for i in range(self.num_buses):
                # Gs and Bs values
                setattr(self, f"Gs{i + 1}_space", tk.DoubleVar(value=Gs_array_loaded[i]))
                setattr(self, f"Bs{i + 1}_space", tk.DoubleVar(value=Bs_array_loaded[i]))

                # PG and QG values
                setattr(self, f"PG{i + 1}_space", tk.DoubleVar(value=self.PG_array_loaded[i]))
                setattr(self, f"QG{i + 1}_space", tk.DoubleVar(value=self.QG_array_loaded[i]))

                # PL and QL values
                setattr(self, f"PL{i + 1}_space", tk.DoubleVar(value=PL_array_loaded[i]))
                setattr(self, f"QL{i + 1}_space", tk.DoubleVar(value=QL_array_loaded[i]))

                # Q_min and Q_max values
                setattr(self, f"Q_min{i + 1}_space", tk.DoubleVar(value=Q_min_array_loaded[i]))
                setattr(self, f"Q_max{i + 1}_space", tk.DoubleVar(value=Q_max_array_loaded[i]))

                # V values
                # Assuming you want to store only the real part of the complex numbers
                setattr(self, f"V{i + 1}_space", tk.DoubleVar(value=V_array_loaded[i].real))

            for i, bus_type in enumerate(bus_type_array_loaded):
                if bus_type == 'Slack':
                    self.slack_bus_option.set(i)  # Assuming slack_bus_option is set to the index of the slack bus
                else:
                    # Assuming for 'PV' bus_options is set to 0 and for 'PQ' it is set to 1
                    option_value = 0 if bus_type == 'PV' else 1
                    self.bus_options[i].set(option_value)

            if self.method == '- PLF-NR' or self.method == '- PLF-NRFD' or self.method == "dyn Config":
                std_dev_array_loaded = Data_Store.central_data_store.get_data('std_dev_array')
                std_0, std_1 = std_dev_array_loaded[0], std_dev_array_loaded[1]
                std_2, std_3 = std_dev_array_loaded[2], std_dev_array_loaded[3]

                self.std_dev_PL_space = tk.DoubleVar(value=std_0)
                self.std_dev_QL_space = tk.DoubleVar(value=std_1)
                self.std_dev_PG_space = tk.DoubleVar(value=std_2)
                self.std_dev_QG_space = tk.DoubleVar(value=std_3)
        self.canvas.bind('<Configure>', self.Entry_setup)

    def on_mousewheel(self, event):
        # Calculate scroll delta
        delta = -1 * (event.delta // 120)

        # Get the current top position of the canvas view
        current_top = self.canvas.yview()[0]

        # Scroll the canvas only if scrolling down or if not exceeding the initial top position
        if delta > 0 or (delta < 0 and current_top > self.initial_top):
            self.canvas.yview_scroll(delta, "units")

    def Entry_setup(self, event):
        # Delete the previous canvas image
        self.canvas.delete("bg_img")
        self.canvas.delete("text_tag")
        self.canvas.delete("error_message")
        self.canvas.delete("hover_area")

        for i in range(self.num_buses):
            self.canvas.delete(f"PG_tag_{i + 1}")
            self.canvas.delete(f"QG_tag_{i + 1}")

        new_width = event.width
        new_height = event.height
        resized_image = self.original_image.resize((new_width, new_height))
        photo = ImageTk.PhotoImage(resized_image)

        self.canvas.background = photo  # reference!
        self.canvas.create_image(0, 0, image=photo, anchor='nw', tags="bg_img")

        self.canvas.create_text(2 * self.ref_x, 0.5 * self.ref_y, text=f"{self.method} method", tags="text_tag",
                                font=('Verdana', int(self.f20.cget('size') + self.ref_f), 'bold'), justify="left")

        # Maximum height for scrolling
        threading.Thread(target=self.calculate_max_height, daemon=True).start()

        self.canvas.config(scrollregion=(0, 0, photo.width(), max(self.max_scroll_height, photo.height())))

        threading.Thread(target=self.wheel, daemon=True).start()

        # Store the initial top position
        self.initial_top = self.canvas.yview()[0]
        threading.Thread(target=self.Entry_widgets, daemon=True).start()

    def wheel(self):
        self.canvas.bind_all("<MouseWheel>", self.on_mousewheel)

    def Entry_widgets(self):
        style = ttk.Style()
        Back_btn = ttk.Button(self.canvas, text=" ⟵",
                              command=lambda: self.gui_setup.hide_all_tabs(self.previous_canvas),
                              style="Custom.TButton", width=4)
        self.canvas.create_window(0.10944 * self.ref_x, 0.06156 * self.ref_y, window=Back_btn)

        self.canvas.create_text(0.35 * self.ref_x, 1 * self.ref_y, text="Line", tags="text_tag",
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
        y_start = 1.35  # Starting y-coordinate
        gap = 0.35
        y_coord = y_start
        for i in range(1, self.num_buses):
            for j in range(i + 1, self.num_buses + 1):
                pair = f"{i}_{j}"
                self.canvas.create_text(0.35 * self.ref_x, y_coord * self.ref_y,
                                        text=f"{pair}:", tags="text_tag", justify="center",
                                        font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))
                y_coord += gap

        x_positions = [0.85, 2, 3.15, 4.25]
        prefixes_names = ["R", "X", "jB", "T"]  # Text prefixes

        y_coord = y_start
        for i in range(1, self.num_buses):
            for j in range(i + 1, self.num_buses + 1):
                for x_pos, text_prefix in zip(x_positions, prefixes_names):
                    self.canvas.create_text(x_pos * self.ref_x, y_coord * self.ref_y,
                                            text=f"{text_prefix}", tags="text_tag", justify="center",
                                            font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))
                y_coord += gap
        prefixes = ["R", "X", "B", "T"]  # Text prefixes

        x_positions = [1.4, 2.55, 3.7, 4.75]  # x-coordinates for R, X, and B widgets

        y_coord = y_start
        for i in range(1, self.num_buses):
            for j in range(i + 1, self.num_buses + 1):
                pair = f"{i}_{j}"
                for x_pos, prefix in zip(x_positions, prefixes):
                    variable_name = f"{prefix}{pair}_space"
                    text_variable = getattr(self, variable_name)

                    widget = ttk.Entry(self.canvas, textvariable=text_variable, width=6, justify="center",
                                       font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
                    self.canvas.create_window(x_pos * self.ref_x, y_coord * self.ref_y, window=widget, tags="text_tag")
                y_coord += gap

        base_x, gap = 6.05, 0.85
        texts = ["V", "Ang."]

        for i, text in enumerate(texts):
            x_coord = (base_x + i * gap) * self.ref_x
            self.canvas.create_text(x_coord, 1.7 * self.ref_y, text=text, tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        base_x = 8
        texts = ["P_G", "Q_G", "P_L", "Q_L", "Q min", "Q max", "Gs", "jBs"]

        for i, text in enumerate(texts):
            x_coord = (base_x + i * gap) * self.ref_x
            self.canvas.create_text(x_coord, 1.7 * self.ref_y, text=text, tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        base_x, gap = 14.55, 0.5
        texts = ['Slack', 'PV', 'PQ']

        for i, text in enumerate(texts):
            x_coord = (base_x + i * gap) * self.ref_x
            self.canvas.create_text(x_coord, 1.7 * self.ref_y, text=text, tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        self.canvas.create_text(5.45 * self.ref_x, 1.7 * self.ref_y, text="Bus", tags="text_tag",
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        x_coord, y_start, gap = 5.45, 2.05, 0.35
        for i in range(1, self.num_buses + 1):
            y_coord = y_start + (i - 1) * gap
            self.canvas.create_text(x_coord * self.ref_x, y_coord * self.ref_y, text=f"{i}.", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        PG_x = 8
        properties = [("V", 6.05), ("ang", 6.9), ("PL", 9.7), ("QL", 10.55),
                      ("Q_min", 11.4), ("Q_max", 12.25), ("Gs", 13.1), ("Bs", 13.95)]
        for i in range(1, self.num_buses + 1):
            y_coord = y_start + (i - 1) * gap
            for prop, x_coord in properties:
                variable_name = f"{prop}{i}_space"
                text_variable = getattr(self, variable_name)

                widget = ttk.Entry(self.canvas, textvariable=text_variable, width=6,
                                   justify="center", font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
                self.canvas.create_window(x_coord * self.ref_x, y_coord * self.ref_y, window=widget, tags="text_tag")

        style.configure("Custom_add.TButton",
                        font=('Verdana', int(self.f12.cget('size') + self.ref_f)))
        self.btn_list = []
        for i in range(1, self.num_buses + 1):
            self.gen_num = i
            btn_y = y_start + (i - 1) * gap  # Calculate y position for each button
            btn = ttk.Button(self.canvas, text="+", width=1, style="Custom_add.TButton",
                             command=lambda num=self.gen_num: self.open_Gen_dialog(num))
            self.canvas.create_window((PG_x - 0.55) * self.ref_x, btn_y * self.ref_y, window=btn)
            self.btn_list.append(btn)
            ToolTip(btn, text="add generators", justify="center")

        properties = [("PG", PG_x), ("QG", 8.85)]
        for i in range(1, self.num_buses + 1):
            y_coord = y_start + (i - 1) * gap
            for prop, x_coord in properties:
                tag = f"{prop}_tag_{i}"
                text_content = "0.0"  # Default text content

                # Check if data arrays are not None and access them safely
                if prop == "PG" and self.PG_array_loaded:
                    text_content = str(self.PG_array_loaded[i - 1])
                if prop == "QG" and self.QG_array_loaded:
                    text_content = str(self.QG_array_loaded[i - 1])

                # Create text on canvas with dynamic positioning and content
                self.canvas.create_text(x_coord * self.ref_x, y_coord * self.ref_y,
                                        text=text_content, justify="center", tags=tag,
                                        font=('Verdana', self.f14.cget('size'), 'bold'))

        self.canvas.create_text(6.75 * self.ref_x, 1.35 * self.ref_y, text="tolerance:  1e-", tags="text_tag",
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        self.canvas.create_text(10 * self.ref_x, 1.35 * self.ref_y, text="max iterations", tags="text_tag",
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        widget = ttk.Entry(self.canvas, textvariable=self.tolerance_space,
                           width=6, justify="center",
                           font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
        self.canvas.create_window(7.9 * self.ref_x, 1.35 * self.ref_y, window=widget, tags="text_tag")

        widget = ttk.Entry(self.canvas, textvariable=self.max_iterations_space, width=6, justify="center",
                           font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
        self.canvas.create_window(11.3 * self.ref_x, 1.35 * self.ref_y, window=widget, tags="text_tag")

        if self.method == '- PLF-NR' or self.method == '- PLF-NRFD' or self.method == "dyn Config":

            prefixes = ["SD_PG", "SD_QG", "SD_PL", "SD_QL"]  # Text prefixes
            base_x = 7.05
            gap = 1.70

            for i, text in enumerate(prefixes):
                x_coord = (base_x + i * gap) * self.ref_x
                self.canvas.create_text(x_coord,
                                        1 * self.ref_y,
                                        text=text, tags="text_tag", justify="center",
                                        font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))
            prefixes = ["std_dev_PG", "std_dev_QG", "std_dev_PL", "std_dev_QL"]  # Text prefixes

            x_positions = [7.90, 9.60, 11.30, 13.00]  # x-coordinates for R, X, and B widgets

            for x_pos, prefix in zip(x_positions, prefixes):
                variable_name = f"{prefix}_space"
                text_variable = getattr(self, variable_name)

                widget = ttk.Entry(self.canvas, textvariable=text_variable,
                                   width=6, justify="center", font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
                self.canvas.create_window(x_pos * self.ref_x,
                                          1 * self.ref_y, window=widget, tags="text_tag")

            self.canvas.create_text(14.1 * self.ref_x,
                                    1 * self.ref_y,
                                    text="Simulations", tags="text_tag", justify="center",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))
            num_simulations_widget = ttk.Entry(self.canvas, textvariable=self.num_simulations_space,
                                               width=6, justify="center",
                                               font=('Verdana', int(self.f14.cget('size') + self.ref_f)))

            self.canvas.create_window(15.25 * self.ref_x,
                                      1 * self.ref_y, window=num_simulations_widget, tags="text_tag")

        style.configure('TRadiobutton', relief='sunken',
                        focuscolor=style.lookup('TButton', 'focuscolor'))

        style.configure("Custom.TButton", font=('Verdana', int(self.f14.cget('size') + self.ref_f)))

        self.Run_btn = ttk.Button(self.canvas, text="▶ Run",
                                  command=lambda: self.start_process(), style="Custom.TButton")

        self.canvas.create_window(8.25 * self.ref_x, 0.5 * self.ref_y, window=self.Run_btn)
        ToolTip(self.Run_btn, text="Start Load Flow Calculations", justify="center")

        self.Processing_text = self.canvas.create_text(8.25 * self.ref_x, 0.15 * self.ref_y, text="",
            font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))

        self.save_btn = ttk.Button(self.canvas, text="💾 Save Entries",
                                   command=lambda: self.save_data(), style="Custom.TButton")
        self.canvas.create_window(6.30 * self.ref_x, 0.5 * self.ref_y, window=self.save_btn)
        ToolTip(self.save_btn, text="Save Data for later use", justify="center")

        plot_btn = ttk.Button(self.canvas, text="Visualize Network",
                              command=lambda: self.Validate_entry(True), style="Custom.TButton")
        self.canvas.create_window(4.25 * self.ref_x, 1 * self.ref_y, window=plot_btn)
        ToolTip(plot_btn, text="Check Network connections", justify="center")

        self.tip_setup()

    def open_Gen_dialog(self, bus_num):
        last_P = np.float64(self.canvas.itemcget(f"PG_tag_{bus_num}", "text"))
        last_Q = np.float64(self.canvas.itemcget(f"QG_tag_{bus_num}", "text"))
        dialog = Gen_dialog(self.canvas, bus_num, last_P, last_Q, f"Generators for Bus {bus_num}")
        total_PG_sub = dialog.total_PG_sub
        total_QG_sub = dialog.total_QG_sub
        text_1 = str(total_PG_sub) if total_PG_sub is not None else last_P
        text_2 = str(total_QG_sub) if total_QG_sub is not None else last_Q

        self.canvas.itemconfig(f"PG_tag_{bus_num}", text=text_1)
        self.canvas.itemconfig(f"QG_tag_{bus_num}", text=text_2)

    def setup_radio_buttons(self):
        # Slack option setup
        for i in range(self.num_buses):
            slack_widget = ttk.Radiobutton(
                self.canvas, variable=self.slack_bus_option, value=i,
                command=self.make_slack_command(i), takefocus=0)
            self.canvas.create_window(
                14.55 * self.ref_x,
                (2.05 + i * 0.35) * self.ref_y, window=slack_widget)
            self.slack_widgets.append(slack_widget)

        # PV - PQ options setup
        for i in range(self.num_buses):
            pv_widget = ttk.Radiobutton(
                self.canvas, variable=self.bus_options[i], value=0,
                command=self.make_pv_pq_command(i, 0), takefocus=0)
            pq_widget = ttk.Radiobutton(
                self.canvas, variable=self.bus_options[i], value=1,
                command=self.make_pv_pq_command(i, 1), takefocus=0)

            self.canvas.create_window(
                (15.05 + 0 * 0.5) * self.ref_x,
                (2.05 + i * 0.35) * self.ref_y,
                window=pv_widget)

            self.canvas.create_window(
                (15.05 + 1 * 0.5) * self.ref_x,
                (2.05 + i * 0.35) * self.ref_y,
                window=pq_widget)

            self.PV_widgets.append(pv_widget)
            self.PQ_widgets.append(pq_widget)

    def make_slack_command(self, index):
        def command():
            # Deactivate the corresponding PV and PQ radio buttons when a slack button is selected
            for i in range(self.num_buses):
                # Here, we need to directly call the config method on the widgets
                self.PV_widgets[i].config(state='disabled' if i == index else 'normal')
                self.PQ_widgets[i].config(state='disabled' if i == index else 'normal')

            # Update the variable to reflect the slack selection
            self.slack_bus_option.set(index)

        return command

    def make_pv_pq_command(self, i, value):
        def command():
            self.pv_pq_option_selected(i, value)

        return command

    def pv_pq_option_selected(self, row, col):
        # Only allow selection if the corresponding slack button is not selected
        if self.slack_bus_option.get() != row:
            self.bus_options[row].set(col)
        else:
            # Deselect if the corresponding slack button is selected
            self.bus_options[row].set(-1)
            # Manually update the appearance of the radio buttons to show they are not selected
            # Ensure you call deselect on the widget itself, not on the canvas item
            self.PV_widgets[row].deselect()
            self.PQ_widgets[row].deselect()

    def check_input_error(self, attribute_name):
        try:
            float(getattr(self, attribute_name).get())
            return True
        except ValueError:
            # Catch ValueError if float conversion fails due to incorrect string formats (non-numeric)
            return False
        except tk.TclError:
            # Catch TclError if tkinter variable handling fails (e.g., bad string passed to DoubleVar)
            return False

    def update_Processing_text(self, show=True):
        if show:
            # Initialize Processing_dots if it's None.
            if self.Processing_dots is None:
                self.Processing_dots = 0

            # Prepare the text based on the current count of dots.
            Processing_message = "Processing" + "." * self.Processing_dots
            self.canvas.itemconfig(self.Processing_text, text=Processing_message)

            # Update the dots count, cycling back after 3.
            self.Processing_dots = (self.Processing_dots + 1) % 4
            if self.thread_error_end:
                self.Run_btn.config(state='normal')
                self.save_btn.config(state='normal')
                self.canvas.itemconfig(self.Processing_text, text="")
                if self.thread_end:
                    self.display_error_message("")
                    self.go_to_result_class()

            else:
                # Schedule the next update
                self.Processing_after = self.canvas.after(500, self.update_Processing_text, True)
        else:
            # When showing is False, stop the updates and clear the text.
            if self.Processing_after is not None:
                self.canvas.after_cancel(self.Processing_after)
                self.Processing_after = self.Processing_dots = None
            self.canvas.itemconfig(self.Processing_text, text="")

    def start_process(self):
        self.thread_end = False
        self.thread_error_end = False
        self.display_error_message("")
        self.Run_btn.config(state='disabled')
        self.save_btn.config(state='disabled')
        self.update_Processing_text()
        threading.Thread(target=self.Validate_entry, args=(None,)).start()

    def Validate_entry(self, case):
        R_check, X_check, B_check, T_check = {}, {}, {}, {}
        Gs_check = [0.0] * self.num_buses
        Bs_check = [0.0] * self.num_buses
        bus_check = [0.0] * self.num_buses
        V_check = [0.0] * self.num_buses
        V_zero_check = [0.0] * self.num_buses
        Ang_check = [0.0] * self.num_buses
        PL_check = [0.0] * self.num_buses
        QL_check = [0.0] * self.num_buses
        Q_min_check = [0.0] * self.num_buses
        Q_max_check = [0.0] * self.num_buses

        # R, X, B, Z, Y:
        # Nested loops to generate pairs and calculate Z values
        for i in range(1, self.num_buses + 1):
            for j in range(i + 1, self.num_buses + 1):
                pair = f"{i}_{j}"
                # Retrieve R, X, and B values for each pair
                R_input = self.check_input_error(f"R{i}_{j}_space")
                if R_input:
                    self.R_values[pair] = np.abs(np.float64(getattr(self, f"R{i}_{j}_space").get()))
                    R_check[pair] = True
                else:
                    R_check[pair] = False

                X_input = self.check_input_error(f"X{i}_{j}_space")
                if X_input:
                    self.X_values[pair] = np.float64(getattr(self, f"X{i}_{j}_space").get())
                    X_check[pair] = True
                else:
                    X_check[pair] = False

                B_input = self.check_input_error(f"B{i}_{j}_space")
                if B_input:
                    self.B_values[pair] = np.float64(getattr(self, f"B{i}_{j}_space").get())
                    B_check[pair] = True
                else:
                    B_check[pair] = False

                T_input = self.check_input_error(f"T{i}_{j}_space")
                if T_input:
                    self.T_values[pair] = np.abs(np.float64(getattr(self, f"T{i}_{j}_space").get()))
                    T_check[pair] = True
                else:
                    T_check[pair] = False

        # Add shunt element admittances to the Y-bus matrix
        for i in range(self.num_buses):
            Gs_input = self.check_input_error(f'Gs{i + 1}_space')
            if Gs_input:
                self.Gs_array[i] = np.float64(getattr(self, f'Gs{i + 1}_space').get())
                Gs_check[i] = True

            else:
                Gs_check[i] = False
            Bs_input = self.check_input_error(f'Bs{i + 1}_space')
            if Bs_input:
                self.Bs_array[i] = np.float64(getattr(self, f'Bs{i + 1}_space').get())
                Bs_check[i] = True
            else:
                Bs_check[i] = False

        # Bus type:
        selected_slack_index = self.slack_bus_option.get()
        if selected_slack_index != -1:
            self.slack_bus_num = selected_slack_index + 1

        # Set bus types based on options
        for i in range(self.num_buses):
            if selected_slack_index != -1 and i + 1 == self.slack_bus_num:
                self.bus_type_array[i] = 'Slack'
                bus_check[selected_slack_index] = True

            else:
                if i == selected_slack_index:
                    continue
                bus_option = self.bus_options[i].get()
                if bus_option == 0:
                    self.bus_type_array[i] = 'PV'
                    bus_check[i] = True
                elif bus_option == 1:
                    self.bus_type_array[i] = 'PQ'
                    bus_check[i] = True
                else:
                    bus_check[i] = False

        for i in range(self.num_buses):
            V_input = self.check_input_error(f'V{i + 1}_space')
            if V_input:
                V_zero_check[i] = np.float64(getattr(self, f'V{i + 1}_space').get())
                V_check[i] = True
            else:
                V_check[i] = False

            ang_input = self.check_input_error(f'ang{i + 1}_space')
            if ang_input:
                Ang_check[i] = True
            else:
                Ang_check[i] = False

        # P, Q Power values:
        for i in range(self.num_buses):
            self.PG_array[i] = np.float64(self.canvas.itemcget(f"PG_tag_{i + 1}", "text"))
            self.QG_array[i] = np.float64(self.canvas.itemcget(f"QG_tag_{i + 1}", "text"))

            PL_input = self.check_input_error(f'PL{i + 1}_space')
            if PL_input:
                self.PL_array[i] = np.abs(np.float64(getattr(self, f'PL{i + 1}_space').get()))
                PL_check[i] = True
            else:
                PL_check[i] = False

            QL_input = self.check_input_error(f'QL{i + 1}_space')
            if QL_input:
                self.QL_array[i] = np.abs(np.float64(getattr(self, f'QL{i + 1}_space').get()))
                QL_check[i] = True
            else:
                QL_check[i] = False

            Q_min_input = self.check_input_error(f'Q_min{i + 1}_space')
            if Q_min_input:
                self.Q_min_array[i] = np.abs(np.float64(getattr(self, f'Q_min{i + 1}_space').get()))
                Q_min_check[i] = True
            else:
                Q_min_check[i] = False

            Q_max_input = self.check_input_error(f'Q_max{i + 1}_space')
            if Q_max_input:
                self.Q_max_array[i] = np.abs(np.float64(getattr(self, f'Q_max{i + 1}_space').get()))
                Q_max_check[i] = True
            else:
                Q_max_check[i] = False

        # List of attribute suffixes
        suffixes = ['PL', 'QL', 'PG', 'QG']

        # Iterate over the suffixes using their index and value
        for i, suffix in enumerate(suffixes):
            # Dynamic attribute name creation
            attr_name = f'std_dev_{suffix}_space'

            # Check if the input error method returns True
            if self.check_input_error(attr_name) and float(getattr(self, attr_name).get()) > 0:
                # Retrieve the value, convert it to float, and assign to the respective index
                self.std_dev_array[i] = float(getattr(self, attr_name).get())
            else:
                self.display_error_message(f"Invalid std_dev_{suffix}")
                self.thread_error_end = True
                return

        if self.check_input_error(f'tolerance_space') and int(self.tolerance_space.get()) > 0:
            self.tolerance = 10 ** -int(self.tolerance_space.get())
        else:
            self.display_error_message("Invalid tolerance")
            self.thread_error_end = True
            return

        if self.check_input_error(f'max_iterations_space') and int(self.max_iterations_space.get()) > 0:
            self.max_iterations = int(self.max_iterations_space.get())
        else:
            self.display_error_message("Invalid max iterations")
            self.thread_error_end = True

            return

        dict_validations = {"R": R_check, "X": X_check, "B": B_check, "T": T_check}
        list_validations = {
            "V.": list(V_check), "V": list(V_zero_check), "Ang": list(Ang_check),
            "PL": list(PL_check), "QL": list(QL_check), "Q_min": list(Q_min_check),
            "Q_max": list(Q_max_check), "type": list(bus_check)}

        # For dictionary validations
        dict_failed_validations = []
        for name, validation_dict in dict_validations.items():
            failed_keys = [key for key, value in validation_dict.items() if not value]
            if failed_keys:
                dict_failed_validations.extend([f"Line {key} {name} " for key in failed_keys])

        # For list validations
        list_failed_validations = []
        for name, validation_list in list_validations.items():
            if name == "V":  # Special case for Voltage magnitude
                failed_indices = [i for i, value in enumerate(validation_list) if value == 0.0]
            else:
                failed_indices = [i for i, value in enumerate(validation_list) if not value]

            if failed_indices:
                list_failed_validations.extend([f"Bus {i + 1} {name} " for i in failed_indices])

        # Combine failed validations from both dictionaries and lists
        failed_validations = dict_failed_validations + list_failed_validations

        if failed_validations:
            error_message = "Invalid input at: " + "; ".join(failed_validations)
            self.display_error_message(error_message)
            self.thread_error_end = True
        else:
            self.process_entry(case)

    def process_entry(self, case):
        # Calculate Z values
        for i in range(1, self.num_buses + 1):
            for j in range(i + 1, self.num_buses + 1):
                pair = f"{i}_{j}"
                self.Z_values[pair] = self.R_values[pair] + 1j * self.X_values[pair]

        for i in range(self.num_buses):
            V = np.float64(getattr(self, f'V{i + 1}_space').get())
            ang = np.radians(np.float64(getattr(self, f'ang{i + 1}_space').get()))
            self.V_array[i] = V * np.exp(1j * ang)

        Y_instances = Y_bus_class(self.num_buses)

        # Add line admittances to the Y-bus matrix
        for pair, Z in self.Z_values.items():
            i, j = map(int, pair.split('_'))  # Correctly split the pair to extract i and j
            Y = 1 / Z if Z != 0 else 0
            B = self.B_values[pair] if pair in self.B_values else 0
            T = self.T_values[pair] if pair in self.T_values else 0
            Y_instances.add_line_admittance(i, j, Y, B, T)
            print(i, j, Y, B, T)
        for i in range(self.num_buses):
            Gs = self.Gs_array[i]
            Bs = self.Bs_array[i]
            Y_instances.add_shunt_element_admittance(i + 1, Gs, Bs)

        # Get the Y-bus matrix
        Y_bus_matrix = Y_instances.get_Y_bus_matrix()
        Data_Store.central_data_store.update_data('R_values', self.R_values)
        Data_Store.central_data_store.update_data('X_values', self.X_values)
        Data_Store.central_data_store.update_data('B_values', self.B_values)
        Data_Store.central_data_store.update_data('T_values', self.T_values)
        Data_Store.central_data_store.update_data('Gs_array', self.Gs_array)
        Data_Store.central_data_store.update_data('Bs_array', self.Bs_array)
        Data_Store.central_data_store.update_data('Y_bus_matrix', Y_bus_matrix)
        print("Y_bus_matrix", Y_bus_matrix)

        if case:
            self.display_error_message("")
            Network = Plots.PowerNetwork(False)
            Network.visualize_network()
        else:
            self.methods_processing(Y_bus_matrix)
            print(Y_bus_matrix)

    def methods_processing(self, Y_bus_matrix):
        if self.method == '- GS':
            solver = Power_flow.GS_class(self.num_buses, self.V_array, Y_bus_matrix, self.PG_array, self.QG_array,
                                         self.PL_array, self.QL_array, self.Q_min_array, self.Q_max_array,
                                         self.slack_bus_num, self.bus_type_array, self.max_iterations, self.tolerance)

            # Output voltage array
            V_output_array, Q_PV, iteration_result = solver.GS_solve()
            V_mag_output_array = np.abs(V_output_array)
            V_ang_output_array = np.angle(V_output_array, deg=True)

            PG_output_array, QG_output_array = Calc_class.Power_generated(
                self.num_buses, self.bus_type_array, self.slack_bus_num,
                Y_bus_matrix, V_output_array, self.PG_array, self.QG_array)

            # Calculate line currents I_ij_array
            I_ij_array = Calc_class.line_currents(self.num_buses, Y_bus_matrix, V_output_array)

            # Calculate line power flows S_ij_array
            S_ij_array = Calc_class.line_power_flows(self.num_buses, V_output_array, I_ij_array)

            # Calculate line losses
            total_P_loss, total_Q_loss = Calc_class.line_losses(
                PG_output_array, QG_output_array, self.PL_array, self.QL_array)

            is_power_balanced, total_P_generation, total_P_load_loss = Calc_class.check_power_balance(
                self.bus_type_array, self.PG_array, self.PL_array, total_P_loss)
            print("Power Balance Maintained:", is_power_balanced)

            Data_Store.central_data_store.update_data('iteration_result', iteration_result)
            Data_Store.central_data_store.update_data('bus_type_array', self.bus_type_array)
            Data_Store.central_data_store.update_data('V_mag_output_array', V_mag_output_array)
            Data_Store.central_data_store.update_data('V_ang_output_array', V_ang_output_array)
            Data_Store.central_data_store.update_data('PG_output_array', PG_output_array)
            Data_Store.central_data_store.update_data('QG_output_array', QG_output_array)
            # Loads stay without change
            Data_Store.central_data_store.update_data('PL_array', self.PL_array)
            Data_Store.central_data_store.update_data('QL_array', self.QL_array)
            Data_Store.central_data_store.update_data('total_P_loss', total_P_loss)
            Data_Store.central_data_store.update_data('total_Q_loss', total_Q_loss)
            Data_Store.central_data_store.update_data('total_P_generation', total_P_generation)
            Data_Store.central_data_store.update_data('total_P_load_loss', total_P_load_loss)
            Data_Store.central_data_store.update_data('I_ij_array', I_ij_array)
            Data_Store.central_data_store.update_data('S_ij_array', S_ij_array)

        elif self.method == '- NR':
            solver = Power_flow.NR_class(self.num_buses, self.V_array, Y_bus_matrix,
                                         self.PG_array, self.QG_array, self.PL_array, self.QL_array,
                                         self.Q_min_array, self.Q_max_array, self.slack_bus_num,
                                         self.bus_type_array, self.max_iterations, self.tolerance)

            # Output voltage array
            V_output_array, iteration_result = solver.NR_solve()
            V_mag_output_array = np.abs(V_output_array)
            V_ang_output_array = np.angle(V_output_array, deg=True)

            PG_output_array, QG_output_array = Calc_class.Power_generated(
                self.num_buses, self.bus_type_array, self.slack_bus_num,
                Y_bus_matrix, V_output_array, self.PG_array, self.QG_array)

            # Calculate line currents I_ij_array
            I_ij_array = Calc_class.line_currents(self.num_buses, Y_bus_matrix, V_output_array)

            # Calculate line power flows S_ij_array
            S_ij_array = Calc_class.line_power_flows(self.num_buses, V_output_array, I_ij_array)

            # Calculate line losses
            total_P_loss, total_Q_loss = Calc_class.line_losses(
                PG_output_array, QG_output_array, self.PL_array, self.QL_array)

            # Validate results
            is_power_balanced, total_P_generation, total_P_load_loss = Calc_class.check_power_balance(
                self.bus_type_array, self.PG_array, self.PL_array, total_P_loss)

            print("Power Balance Maintained:", is_power_balanced)

            Data_Store.central_data_store.update_data('iteration_result', iteration_result)
            Data_Store.central_data_store.update_data('bus_type_array', self.bus_type_array)
            Data_Store.central_data_store.update_data('V_mag_output_array', V_mag_output_array)
            Data_Store.central_data_store.update_data('V_ang_output_array', V_ang_output_array)
            Data_Store.central_data_store.update_data('PG_output_array', PG_output_array)
            Data_Store.central_data_store.update_data('QG_output_array', QG_output_array)
            # Loads stay without change
            Data_Store.central_data_store.update_data('PL_array', self.PL_array)
            Data_Store.central_data_store.update_data('QL_array', self.QL_array)
            Data_Store.central_data_store.update_data('total_P_loss', total_P_loss)
            Data_Store.central_data_store.update_data('total_Q_loss', total_Q_loss)
            Data_Store.central_data_store.update_data('total_P_generation', total_P_generation)
            Data_Store.central_data_store.update_data('total_P_load_loss', total_P_load_loss)
            Data_Store.central_data_store.update_data('I_ij_array', I_ij_array)
            Data_Store.central_data_store.update_data('S_ij_array', S_ij_array)

        elif self.method == '- NRFD':
            solver = Power_flow.NRFD_class(self.num_buses, self.V_array, Y_bus_matrix, self.PG_array, self.QG_array,
                                           self.PL_array, self.QL_array, self.Q_min_array, self.Q_max_array,
                                           self.slack_bus_num, self.bus_type_array, self.max_iterations, self.tolerance)

            # Output voltage array
            V_output_array, iteration_result = solver.NRFD_solve()
            V_mag_output_array = np.abs(V_output_array)
            V_ang_output_array = np.angle(V_output_array, deg=True)

            PG_output_array, QG_output_array = Calc_class.Power_generated(
                self.num_buses, self.bus_type_array, self.slack_bus_num,
                Y_bus_matrix, V_output_array, self.PG_array, self.QG_array)

            # Calculate line currents I_ij_array
            I_ij_array = Calc_class.line_currents(self.num_buses, Y_bus_matrix, V_output_array)

            # Calculate line power flows S_ij_array
            S_ij_array = Calc_class.line_power_flows(self.num_buses, V_output_array, I_ij_array)

            # Calculate line losses
            total_P_loss, total_Q_loss = Calc_class.line_losses(
                PG_output_array, QG_output_array, self.PL_array, self.QL_array)

            # Validate results
            is_power_balanced, total_P_generation, total_P_load_loss = Calc_class.check_power_balance(
                self.bus_type_array, self.PG_array, self.PL_array, total_P_loss)

            print("Power Balance Maintained:", is_power_balanced)

            Data_Store.central_data_store.update_data('iteration_result', iteration_result)
            Data_Store.central_data_store.update_data('bus_type_array', self.bus_type_array)
            Data_Store.central_data_store.update_data('V_mag_output_array', V_mag_output_array)
            Data_Store.central_data_store.update_data('V_ang_output_array', V_ang_output_array)
            Data_Store.central_data_store.update_data('PG_output_array', PG_output_array)
            Data_Store.central_data_store.update_data('QG_output_array', QG_output_array)
            # Loads stay without change
            Data_Store.central_data_store.update_data('PL_array', self.PL_array)
            Data_Store.central_data_store.update_data('QL_array', self.QL_array)
            Data_Store.central_data_store.update_data('total_P_loss', total_P_loss)
            Data_Store.central_data_store.update_data('total_Q_loss', total_Q_loss)
            Data_Store.central_data_store.update_data('total_P_generation', total_P_generation)
            Data_Store.central_data_store.update_data('total_P_load_loss', total_P_load_loss)
            Data_Store.central_data_store.update_data('I_ij_array', I_ij_array)
            Data_Store.central_data_store.update_data('S_ij_array', S_ij_array)

        elif self.method == '- PLF-NR':
            std_dev_array = np.array([0.1] * 4)
            std_dev_array[0] = float(self.std_dev_PL_space.get())
            std_dev_array[1] = float(self.std_dev_QL_space.get())
            std_dev_array[2] = float(self.std_dev_PG_space.get())
            std_dev_array[3] = float(self.std_dev_QG_space.get())
            num_simulations = self.num_simulations_space.get()

            PLF_solver = PLF_calc.Probabilistic_NR_load_flow(self.num_buses, self.V_array, Y_bus_matrix,
                                                             self.PG_array, self.QG_array, self.PL_array, self.QL_array,
                                                             self.Q_min_array, self.Q_max_array,
                                                             self.slack_bus_num, self.bus_type_array, std_dev_array,
                                                             num_simulations,
                                                             self.max_iterations, self.tolerance)

            # Output PLF:
            all_results, non_convergence_log = PLF_solver.run_probabilistic_load_flow()

            # make a normal NR calculations after for reference:
            solver = Power_flow.NR_class(self.num_buses, self.V_array, Y_bus_matrix,
                                         self.PG_array, self.QG_array, self.PL_array, self.QL_array,
                                         self.Q_min_array, self.Q_max_array, self.slack_bus_num,
                                         self.bus_type_array, self.max_iterations, self.tolerance)

            # Output voltage array
            V_output_array, iteration_result = solver.NR_solve()
            V_mag_output_array = np.abs(V_output_array)
            V_ang_output_array = np.angle(V_output_array, deg=True)

            PG_output_array, QG_output_array = Calc_class.Power_generated(
                self.num_buses, self.bus_type_array, self.slack_bus_num,
                Y_bus_matrix, V_output_array, self.PG_array, self.QG_array)

            # Calculate line currents I_ij_array
            I_ij_array = Calc_class.line_currents(self.num_buses, Y_bus_matrix, V_output_array)

            # Calculate line power flows S_ij_array
            S_ij_array = Calc_class.line_power_flows(self.num_buses, V_output_array, I_ij_array)

            # Calculate line losses
            total_P_loss, total_Q_loss = Calc_class.line_losses(
                PG_output_array, QG_output_array, self.PL_array, self.QL_array)

            # Validate results
            is_power_balanced, total_P_generation, total_P_load_loss = Calc_class.check_power_balance(
                self.bus_type_array, self.PG_array, self.PL_array, total_P_loss)

            print("Power Balance Maintained:", is_power_balanced)

            Data_Store.central_data_store.update_data('all_results', all_results)
            Data_Store.central_data_store.update_data('non_convergence_log', non_convergence_log)
            Data_Store.central_data_store.update_data('num_simulations', num_simulations)
            Data_Store.central_data_store.update_data('slack_bus_num', self.slack_bus_num)
            Data_Store.central_data_store.update_data('iteration_result', iteration_result)
            Data_Store.central_data_store.update_data('bus_type_array', self.bus_type_array)
            Data_Store.central_data_store.update_data('V_mag_output_array', V_mag_output_array)
            Data_Store.central_data_store.update_data('V_ang_output_array', V_ang_output_array)
            Data_Store.central_data_store.update_data('PG_output_array', PG_output_array)
            Data_Store.central_data_store.update_data('QG_output_array', QG_output_array)
            # Loads stay without change
            Data_Store.central_data_store.update_data('PL_array', self.PL_array)
            Data_Store.central_data_store.update_data('QL_array', self.QL_array)
            Data_Store.central_data_store.update_data('total_P_loss', total_P_loss)
            Data_Store.central_data_store.update_data('total_Q_loss', total_Q_loss)
            Data_Store.central_data_store.update_data('total_P_generation', total_P_generation)
            Data_Store.central_data_store.update_data('total_P_load_loss', total_P_load_loss)
            Data_Store.central_data_store.update_data('I_ij_array', I_ij_array)
            Data_Store.central_data_store.update_data('S_ij_array', S_ij_array)

        elif self.method == '- PLF-NRFD':
            std_dev_array = np.array([0.1] * 4)
            std_dev_array[0] = float(self.std_dev_PL_space.get())
            std_dev_array[1] = float(self.std_dev_QL_space.get())
            std_dev_array[2] = float(self.std_dev_PG_space.get())
            std_dev_array[3] = float(self.std_dev_QG_space.get())
            num_simulations = self.num_simulations_space.get()

            PLF_solver = PLF_calc.Probabilistic_NRFD_load_flow(self.num_buses, self.V_array, Y_bus_matrix,
                                                               self.PG_array, self.QG_array, self.PL_array,
                                                               self.QL_array, self.Q_min_array, self.Q_max_array,
                                                               self.slack_bus_num, self.bus_type_array, std_dev_array,
                                                               num_simulations,
                                                               self.max_iterations, self.tolerance)

            # Output PLF:
            all_results, non_convergence_log = PLF_solver.run_probabilistic_load_flow()
            # make a normal NRFD calculations after for reference:
            solver = Power_flow.NRFD_class(self.num_buses, self.V_array, Y_bus_matrix,
                                           self.PG_array, self.QG_array, self.PL_array, self.QL_array,
                                           self.Q_min_array, self.Q_max_array, self.slack_bus_num,
                                           self.bus_type_array, self.max_iterations, self.tolerance)

            # Output voltage array
            V_output_array, iteration_result = solver.NRFD_solve()
            V_mag_output_array = np.abs(V_output_array)
            V_ang_output_array = np.angle(V_output_array, deg=True)

            PG_output_array, QG_output_array = Calc_class.Power_generated(
                self.num_buses, self.bus_type_array, self.slack_bus_num,
                Y_bus_matrix, V_output_array, self.PG_array, self.QG_array)

            # Calculate line currents I_ij_array
            I_ij_array = Calc_class.line_currents(self.num_buses, Y_bus_matrix, V_output_array)

            # Calculate line power flows S_ij_array
            S_ij_array = Calc_class.line_power_flows(self.num_buses, V_output_array, I_ij_array)

            # Calculate line losses
            total_P_loss, total_Q_loss = Calc_class.line_losses(
                PG_output_array, QG_output_array, self.PL_array, self.QL_array)

            # Validate results
            is_power_balanced, total_P_generation, total_P_load_loss = Calc_class.check_power_balance(
                self.bus_type_array, self.PG_array, self.PL_array, total_P_loss)

            print("Power Balance Maintained:", is_power_balanced)

            Data_Store.central_data_store.update_data('all_results', all_results)
            Data_Store.central_data_store.update_data('non_convergence_log', non_convergence_log)
            Data_Store.central_data_store.update_data('num_simulations', num_simulations)
            Data_Store.central_data_store.update_data('slack_bus_num', self.slack_bus_num)
            Data_Store.central_data_store.update_data('iteration_result', iteration_result)
            Data_Store.central_data_store.update_data('bus_type_array', self.bus_type_array)
            Data_Store.central_data_store.update_data('V_mag_output_array', V_mag_output_array)
            Data_Store.central_data_store.update_data('V_ang_output_array', V_ang_output_array)
            Data_Store.central_data_store.update_data('PG_output_array', PG_output_array)
            Data_Store.central_data_store.update_data('QG_output_array', QG_output_array)
            # Loads stay without change
            Data_Store.central_data_store.update_data('PL_array', self.PL_array)
            Data_Store.central_data_store.update_data('QL_array', self.QL_array)
            Data_Store.central_data_store.update_data('total_P_loss', total_P_loss)
            Data_Store.central_data_store.update_data('total_Q_loss', total_Q_loss)
            Data_Store.central_data_store.update_data('total_P_generation', total_P_generation)
            Data_Store.central_data_store.update_data('total_P_load_loss', total_P_load_loss)
            Data_Store.central_data_store.update_data('I_ij_array', I_ij_array)
            Data_Store.central_data_store.update_data('S_ij_array', S_ij_array)
        self.thread_end = True
        self.thread_error_end = True

    def go_to_result_class(self):
        self.gui_setup.add_result_tab()
        # Run next class:
        Result_class(self.gui_setup.Result_tab, self.gui_setup.data_entry_tab)

    def display_error_message(self, message):
        self.canvas.delete("error_text")
        Dark_red = '#8b0000'
        self.canvas.create_text(12 * self.ref_x, 0.5 * self.ref_y, text=message, fill=Dark_red, tags="error_text",
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))

    def calculate_max_height(self, base_num_buses=8, base_max_height=1000):
        if self.num_buses < base_num_buses:
            self.max_scroll_height = 0
        elif self.num_buses == base_num_buses:
            self.max_scroll_height = base_max_height
        else:
            # Calculate the number of pairs for the base and new number of buses
            base_pairs = base_num_buses * (base_num_buses - 1) / 2
            new_pairs = self.num_buses * (self.num_buses - 1) / 2

            # Calculate the ratio of new pairs to base pairs
            pair_ratio = new_pairs / base_pairs

            # Calculate the new max height based on the pair ratio
            self.max_scroll_height = int(base_max_height * pair_ratio)

    def save_data(self):
        for i in range(1, self.num_buses + 1):
            for j in range(i + 1, self.num_buses + 1):
                pair = f"{i}_{j}"
                # Retrieve R, X, and B values for each pair
                self.R_values[pair] = np.abs(np.float64(getattr(self, f"R{i}_{j}_space").get()))
                self.X_values[pair] = np.float64(getattr(self, f"X{i}_{j}_space").get())
                self.B_values[pair] = np.float64(getattr(self, f"B{i}_{j}_space").get())
                self.T_values[pair] = np.abs(np.float64(getattr(self, f"T{i}_{j}_space").get()))

        for i in range(self.num_buses):
            self.Gs_array[i] = np.float64(getattr(self, f'Gs{i + 1}_space').get())
            self.Bs_array[i] = np.float64(getattr(self, f'Bs{i + 1}_space').get())

        # Bus type:
        selected_slack_index = self.slack_bus_option.get()
        self.slack_bus_num = selected_slack_index + 1 if selected_slack_index != -1 else None
        # Set bus types based on options
        for i in range(self.num_buses):
            if self.slack_bus_num is not None and i + 1 == self.slack_bus_num:
                self.bus_type_array[i] = 'Slack'
            else:
                bus_option = self.bus_options[i].get()
                if bus_option == 0:
                    self.bus_type_array[i] = 'PV'
                elif bus_option == 1:
                    self.bus_type_array[i] = 'PQ'

        # Voltage mag and Angle:
        for i in range(self.num_buses):
            V = np.float64(getattr(self, f'V{i + 1}_space').get())
            ang = np.radians(np.float64(getattr(self, f'ang{i + 1}_space').get()))
            self.V_array[i] = V * np.exp(1j * ang)

        # P, Q Power values:
        for i in range(self.num_buses):
            self.PG_array[i] = np.float64(self.canvas.itemcget(f"PG_tag_{i + 1}", "text"))
            self.QG_array[i] = np.float64(self.canvas.itemcget(f"QG_tag_{i + 1}", "text"))
            self.PL_array[i] = np.abs(np.float64(getattr(self, f'PL{i + 1}_space').get()))
            self.QL_array[i] = np.abs(np.float64(getattr(self, f'QL{i + 1}_space').get()))
            self.Q_min_array[i] = np.abs(np.float64(getattr(self, f'Q_min{i + 1}_space').get()))
            self.Q_max_array[i] = np.abs(np.float64(getattr(self, f'Q_max{i + 1}_space').get()))

        self.std_dev_array[0] = float(self.std_dev_PL_space.get())
        self.std_dev_array[1] = float(self.std_dev_QL_space.get())
        self.std_dev_array[2] = float(self.std_dev_PG_space.get())
        self.std_dev_array[3] = float(self.std_dev_QG_space.get())
        current_date = datetime.date.today()
        data = (
            f"Date: {current_date}\nnum_buses = {self.num_buses}"
            f"\nR_values = {self.R_values}\nX_values = {self.X_values}"
            f"\nB_values = {self.B_values}\nT_values = {self.T_values}"
            f"\nV_array = {str(list(self.V_array))}\nbus_type_array = {str(list(self.bus_type_array))}"
            f"\nPG_array = {str(list(self.PG_array))}\nQG_array = {str(list(self.QG_array))}"
            f"\nPL_array = {str(list(self.PL_array))}\nQL_array = {str(list(self.QL_array))}"
            f"\nQ_min_array = {str(list(self.Q_min_array))}\nQ_max_array = {str(list(self.Q_max_array))}"
            f"\nGs_array = {str(list(self.Gs_array))}\nBs_array = {str(list(self.Bs_array))}"
            f"\nstd_dev_array = {str(list(self.std_dev_array))}")

        for i in range(1, self.num_buses + 1):
            PG_sub_value = Data_Store.central_data_store.get_data(f'PG_sub_{i}') or [0.0]
            QG_sub_value = Data_Store.central_data_store.get_data(f'QG_sub_{i}') or [0.0]

            # Adjust lengths if one is not all zeros and the other is either None or all zeros
            if not all(p == 0 for p in PG_sub_value) or not all(q == 0 for q in QG_sub_value):
                max_length = max(len(PG_sub_value), len(QG_sub_value))
                PG_sub_value = PG_sub_value if not all(p == 0 for p in PG_sub_value) else [0.0] * max_length
                QG_sub_value = QG_sub_value if not all(q == 0 for q in QG_sub_value) else [0.0] * max_length

                data += f"\nPG_sub_{i} = {PG_sub_value}"
                data += f"\nQG_sub_{i} = {QG_sub_value}"

        base_filename = f'{self.num_buses}_Buses_saved_{current_date}'
        filename = base_filename + ".txt"

        full_path = os.path.join('Saved data', filename)
        # Check if the file exists and find a new file name if necessary
        counter = 1
        while os.path.exists(full_path):
            filename = f"{base_filename} ({counter}).txt"
            full_path = os.path.join('Saved data', filename)
            counter += 1

        with open(full_path, 'w') as file:
            file.write(data)

        directory = os.path.dirname(os.path.abspath(full_path))
        subprocess.Popen(f'explorer "{directory}"')

    def tip_setup(self, size=(441, 287)):
        # Define patterns and positions directly in the class
        patterns = {
            'hover1': ('Generators.*', (8 * self.ref_x, 1.7 * self.ref_y)),
            'hover2': ('transmission lines.*', (0.35 * self.ref_x, 1 * self.ref_y)),
            'hover3': ('Busbars.*', (5.45 * self.ref_x, 1.7 * self.ref_y)),
        }

        self.image_info = {}
        for tag, (pattern, anchor) in patterns.items():
            full_pattern = os.path.join("config", pattern)
            for filepath in glob.glob(full_pattern):
                if os.path.isfile(filepath):
                    self.image_info[tag] = (filepath, anchor)
                    break  # Stop after the first match

        self.size = size
        self.tip_windows = {}  # Dictionary to keep track of tip windows

        # Create hover areas and bind events
        for tag, (image_path, anchor_point) in self.image_info.items():
            x, y = anchor_point
            self.canvas.create_text(x, y, text="\t", tags=tag)
            self.canvas.tag_bind(tag, "<Enter>", lambda e, t=tag: self.show_tip(t, e))
            self.canvas.tag_bind(tag, "<Leave>", lambda e, t=tag: self.hide_tip(t))

    def show_tip(self, tag, event):
        # Hide any existing tip window for this tag
        self.hide_tip(tag)

        # Get the anchor point and image path for the current tag
        image_path, anchor_point = self.image_info[tag]
        x = event.x_root + 20
        y = event.y_root + 20

        # Create a toplevel window at the desired position
        self.tip_windows[tag] = tk.Toplevel(self.canvas)
        self.tip_windows[tag].wm_overrideredirect(True)
        self.tip_windows[tag].wm_geometry(f"+{x}+{y}")

        # Load and resize the image
        image = Image.open(image_path)
        image = image.resize(self.size)
        photo = ImageTk.PhotoImage(image)

        # Create a label with the image and pack
        label = tk.Label(self.tip_windows[tag], image=photo)
        label.image = photo  # Keep a reference to the image
        label.pack()

    def hide_tip(self, tag):
        if tag in self.tip_windows and self.tip_windows[tag]:
            self.tip_windows[tag].destroy()
            self.tip_windows[tag] = None


class Result_class:
    def __init__(self, canvas, previous_canvas):
        self.R_Zero_values, self.X_Zero_values, self.assume_Z_widget, self.Z_zero_btn = None, None, None, None
        self.Z_product, self.Save_switch, self.Round_value = 3, 0, 3
        self.selected_directory, self.fault_bus, self.fault_bus, self.label = None, None, None, None
        self.doc = Document()
        self.I_sc_bus, self.SC_type, self.SC_text, self.SC_label, self.notebook = None, None, None, None, None
        self.V_output_array, self.S_ij_fault_array, self.I_ij_fault_array, self.result_frame = None, None, None, None
        self.V_ang_fault_array, self.V_mag_fault_array, self.SC_frame = None, None, None
        self.fault_type_list = ["L-L-L fault", "L-L-G fault", "L-L fault", "L-G fault"]
        self.fault_bus_space = tk.IntVar(value=1)
        self.assume_Z_space = tk.DoubleVar(value=self.Z_product)
        self.Z_zero_option = tk.IntVar()
        self.Z_zero_option.set(1)  # Set default to the first radio button

        self.canvas, self.previous_canvas = canvas, previous_canvas
        self.canvas = tk.Canvas(self.canvas)
        self.canvas.place(relwidth=1, relheight=1)
        self.app = Data_Store.central_data_store.get_data('app')
        self.gui_setup = Data_Store.central_data_store.get_data('gui_setup')
        self.original_image = Data_Store.central_data_store.get_data('original_image')
        self.fonts = Data_Store.central_data_store.get_data('fonts')
        self.ref_x = Data_Store.central_data_store.get_data('ref_x')
        self.ref_y = Data_Store.central_data_store.get_data('ref_y')
        self.ref_f = Data_Store.central_data_store.get_data('ref_f')
        self.num_buses = Data_Store.central_data_store.get_data('num_buses')
        self.method = Data_Store.central_data_store.get_data('method')
        self.iteration_result = Data_Store.central_data_store.get_data('iteration_result')
        self.bus_type_array = Data_Store.central_data_store.get_data('bus_type_array')

        self.V_mag_output_array = Data_Store.central_data_store.get_data('V_mag_output_array')
        self.V_ang_output_array = Data_Store.central_data_store.get_data('V_ang_output_array')

        self.PG_output_array = Data_Store.central_data_store.get_data('PG_output_array')
        self.QG_output_array = Data_Store.central_data_store.get_data('QG_output_array')

        self.PL_array = Data_Store.central_data_store.get_data('PL_array')
        self.QL_array = Data_Store.central_data_store.get_data('QL_array')

        self.total_P_loss = Data_Store.central_data_store.get_data('total_P_loss')
        self.total_Q_loss = Data_Store.central_data_store.get_data('total_Q_loss')

        self.total_P_generation = Data_Store.central_data_store.get_data('total_P_generation')
        self.total_P_load_loss = Data_Store.central_data_store.get_data('total_P_load_loss')

        self.I_ij_array = Data_Store.central_data_store.get_data('I_ij_array')
        self.S_ij_array = Data_Store.central_data_store.get_data('S_ij_array')
        self.S_ij_losses_array = Data_Store.central_data_store.get_data('S_ij_losses_array')

        self.R_values = Data_Store.central_data_store.get_data('R_values')
        self.X_values = Data_Store.central_data_store.get_data('X_values')
        self.B_values = Data_Store.central_data_store.get_data('B_values')
        self.T_values = Data_Store.central_data_store.get_data('T_values')
        self.Gs_array = Data_Store.central_data_store.get_data('Gs_array')
        self.Bs_array = Data_Store.central_data_store.get_data('Bs_array')
        self.Y_bus_matrix = Data_Store.central_data_store.get_data('Y_bus_matrix')

        self.f12, self.f14, self.f20 = self.fonts[1], self.fonts[2], self.fonts[4]

        self.canvas.bind('<Configure>', self.Result_setup)

    def Result_setup(self, event):
        # Delete the previous canvas image
        self.canvas.delete("bg_img")
        self.canvas.delete("text_tag")
        self.canvas.delete("error_message")
        new_width = event.width
        new_height = event.height
        resized_image = self.original_image.resize((new_width, new_height))
        photo = ImageTk.PhotoImage(resized_image)

        self.canvas.background = photo  # reference!
        self.canvas.create_image(0, 0, image=photo, anchor='nw', tags="bg_img")

        Back_btn = ttk.Button(self.canvas, text=" ⟵",
                              command=lambda: self.gui_setup.hide_all_tabs(self.previous_canvas),
                              style="Custom.TButton", width=4)
        self.canvas.create_window(0.10944 * self.ref_x, 0.06156 * self.ref_y, window=Back_btn)

        self.canvas.create_text(2 * self.ref_x, 0.50 * self.ref_y,
                                text=f"{self.method} Result", tags="text_tag",
                                font=('Verdana', int(self.f20.cget('size') + self.ref_f), 'bold'), justify="left")

        self.canvas.create_text(12 * self.ref_x, 0.50 * self.ref_y,
                                text=f"{self.iteration_result}", tags="text_tag",
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        base_x, gap = 8.50, 1
        texts = ["V", "Ang", "P_G", "Q_G", "P_L", "Q_L", "Bus Type"]

        for i, result_text in enumerate(texts):
            x_coord = (base_x + i * gap) * self.ref_x
            self.canvas.create_text(x_coord, 1 * self.ref_y, text=result_text, tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        # Initial x-coordinate and spacing
        x_start, x_spacing = 8.50, 1
        y_start = 1.35  # Starting y-coordinate
        y_increment = 0.35  # Increment for y-coordinate

        # Generate x-coordinates based on the start and spacing
        x_coord = [x_start + j * x_spacing for j in range(7)]

        for i in range(self.num_buses):
            y_coord = y_start + i * y_increment
            # Creating text for each value in the row corresponding to bus i
            self.canvas.create_text(x_coord[0] * self.ref_x, y_coord * self.ref_y,
                                    text=f"{np.round(self.V_mag_output_array[i], 3)}", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
            self.canvas.create_text(x_coord[1] * self.ref_x, y_coord * self.ref_y,
                                    text=f"{np.round(self.V_ang_output_array[i], 3)}°", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
            self.canvas.create_text(x_coord[2] * self.ref_x, y_coord * self.ref_y,
                                    text=f"{np.round(self.PG_output_array[i], 3)}", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
            self.canvas.create_text(x_coord[3] * self.ref_x, y_coord * self.ref_y,
                                    text=f"{np.round(self.QG_output_array[i], 3)}", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
            self.canvas.create_text(x_coord[4] * self.ref_x, y_coord * self.ref_y,
                                    text=f"{np.round(self.PL_array[i], 3)}", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
            self.canvas.create_text(x_coord[5] * self.ref_x, y_coord * self.ref_y,
                                    text=f"{np.round(self.QL_array[i], 3)}", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
            self.canvas.create_text(x_coord[6] * self.ref_x, y_coord * self.ref_y,
                                    text=f"{self.bus_type_array[i]}", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
        x_coord, y_start, y_increment = 7.50, 1.35, 0.35

        for i in range(1, self.num_buses + 1):
            y_coord = y_start + (i - 1) * y_increment
            self.canvas.create_text(x_coord * self.ref_x, y_coord * self.ref_y, text=f"bus {i}", tags="text_tag",
                                    font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        style = ttk.Style()
        style.configure("Custom.TButton", font=('Verdana', int(self.f14.cget('size') + self.ref_f)))

        Export_btn = ttk.Button(self.canvas, text="Export (.docx)", command=lambda: self.pre_save_doc(),
                                style="Custom.TButton")
        self.canvas.create_window(8.25 * self.ref_x, 0.50 * self.ref_y, window=Export_btn)
        self.label = tk.LabelFrame(self.canvas, bg='#528dab', fg='White')
        self.canvas.create_window(3.50 * self.ref_x, 4 * self.ref_y, window=self.label)

        load_gen_btn = ttk.Button(self.canvas, text="Gen/Load Plot", command=lambda: self.plot_load_gen(),
                                  style="Custom.TButton")
        self.canvas.create_window(1.25 * self.ref_x, 7.50 * self.ref_y, window=load_gen_btn)

        V_profile_btn = ttk.Button(self.canvas, text="V Profile Plot",
                                   command=lambda: self.plot_Voltage_profile(), style="Custom.TButton")
        self.canvas.create_window(2.75 * self.ref_x, 7.50 * self.ref_y, window=V_profile_btn)

        PF_btn = ttk.Button(self.canvas, text="PF Plot",
                            command=lambda: self.plot_PF(), style="Custom.TButton")
        self.canvas.create_window(4.25 * self.ref_x, 7.50 * self.ref_y, window=PF_btn)

        Network_btn = ttk.Button(self.canvas, text="Visualize Grid",
                                 command=lambda: self.plot_Network(), style="Custom.TButton")
        self.canvas.create_window(5.75 * self.ref_x, 7.50 * self.ref_y, window=Network_btn)
        if self.method == "- PLF-NR" or self.method == "- PLF-NRFD" or self.method == "dyn Config":
            Probabilistic_btn = ttk.Button(self.canvas, text="Probabilistic Load Flow Result",
                                           command=lambda: self.PLF_window(), style="Custom.TButton")
            self.canvas.create_window(3.48 * self.ref_x, 8 * self.ref_y, window=Probabilistic_btn)
            if self.method == "dyn Config":
                Probabilistic_btn.config(state='disabled')
                Export_btn.config(state='disabled')
                load_gen_btn.config(state='disabled')
                V_profile_btn.config(state='disabled')
                PF_btn.config(state='disabled')
                Network_btn.config(state='disabled')

        self.notebook = ttk.Notebook(self.label, width=500, height=550)
        self.notebook.pack(fill="both", expand=True)
        self.result_frame = ttk.Frame(self.notebook)
        self.SC_frame = ttk.Frame(self.notebook)
        self.result_frame.pack(fill='both', expand=True)
        self.SC_frame.pack(fill='both', expand=True)

        self.notebook.add(self.result_frame, text=f'Power Flow')
        self.notebook.add(self.SC_frame, text=f'Short Circuit')
        result_canvas = tk.Canvas(self.result_frame, bg='#8CA6B8')
        SC_canvas = tk.Canvas(self.SC_frame, bg='#C47B7B')
        result_canvas.pack(fill='both', expand=True)
        SC_canvas.pack(fill='both', expand=True)

        result_text = tk.Text(result_canvas, bg='#8CA6B8', fg='black',
                              font=('Verdana', int(self.f12.cget('size') + self.ref_f), 'bold'))
        result_text.pack(fill='both', expand=True)
        result_text.delete("1.0", tk.END)

        result_text.insert(tk.END, f"Total P losses          :{np.round(self.total_P_loss, 3)}\n")
        result_text.insert(tk.END, f"Total Q losses          :{np.round(self.total_Q_loss, 3)}\n")

        result_text.insert(tk.END, f"Total P generation      :{np.round(self.total_P_generation, 3)}\n")
        result_text.insert(tk.END, f"Total P loads & losses  :{np.round(self.total_P_load_loss, 3)}\n")

        I_ij_array_mag = np.abs(self.I_ij_array)
        I_ij_array_ang = np.angle(self.I_ij_array, deg=True)

        result_text.insert(tk.END, "\nLine\tLine Currents [p.u.]\n\n")
        for i in range(self.num_buses):
            for j in range(self.num_buses):
                if I_ij_array_mag[i][j] != 0 and I_ij_array_ang[i][j] != 0:  # Check if not zero
                    # Access each element with I_ij_array[i][j]
                    element_1 = np.round(I_ij_array_mag[i][j], 3)
                    element_2 = np.round(I_ij_array_ang[i][j], 3)

                    # insert the values in the text widget
                    result_text.insert(tk.END, f"{i + 1}-{j + 1}:\t{element_1} ∠{element_2}°\n")

        result_text.insert(tk.END, "\nLine\tLine Power Flows [p.u.]\n\n")

        S_ij_array_mag = np.abs(self.S_ij_array)
        S_ij_array_ang = np.angle(self.S_ij_array, deg=True)

        for i in range(self.num_buses):
            for j in range(self.num_buses):
                if S_ij_array_mag[i][j] != 0 and S_ij_array_ang[i][j] != 0:  # Check if not zero
                    # Access each element with S_ij_array[i][j]
                    element_1 = np.round(S_ij_array_mag[i][j], 3)
                    element_2 = np.round(S_ij_array_ang[i][j], 3)
                    # insert the values in the text widget
                    result_text.insert(tk.END, f"{i + 1}-{j + 1}:\t{element_1} ∠{element_2}°\n")

        result_text["state"] = "disabled"

        SC_canvas.create_text(120, 25, text="Identify Faulted Bus:", tags="text_tag",
                              font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
        SC_canvas.create_text(120, 75, text="Fault Type:", tags="text_tag",
                              font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
        fault_bus_widget = ttk.Entry(self.canvas, textvariable=self.fault_bus_space, width=3,
                                     justify="center", font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
        SC_canvas.create_window(250, 28, window=fault_bus_widget, tags="text_tag")

        LLL_btn = ttk.Button(self.canvas, text="L-L-L fault",
                             command=lambda: self.SC_calc_command("L-L-L fault"), style="Custom.TButton")
        SC_canvas.create_window(250, 75, window=LLL_btn)

        LLG_btn = ttk.Button(self.canvas, text="L-L-G fault",
                             command=lambda: self.SC_calc_command("L-L-G fault"), style="Custom.TButton")
        SC_canvas.create_window(250, 115, window=LLG_btn)

        LL_btn = ttk.Button(self.canvas, text="L-L fault",
                            command=lambda: self.SC_calc_command("L-L fault"), style="Custom.TButton")
        SC_canvas.create_window(400, 75, window=LL_btn)

        LG_btn = ttk.Button(self.canvas, text="L-G fault",
                            command=lambda: self.SC_calc_command("L-G fault"), style="Custom.TButton")
        SC_canvas.create_window(400, 115, window=LG_btn)

        style = ttk.Style()
        # Create a custom style that inherits from 'TRadiobutton'
        style.configure('CustomTRadiobutton.TRadiobutton', background='#C47B7B',
                        font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")
        radio1 = ttk.Radiobutton(SC_canvas, variable=self.Z_zero_option, value=1, takefocus=0,
                                 style='CustomTRadiobutton.TRadiobutton',
                                 text="Simplified", command=self.on_radio_change)
        radio2 = ttk.Radiobutton(SC_canvas, variable=self.Z_zero_option, value=2, takefocus=0,
                                 style='CustomTRadiobutton.TRadiobutton',
                                 text="Advanced", command=self.on_radio_change)
        SC_canvas.create_window(240, 150, window=radio1, tags="text_tag")
        SC_canvas.create_window(240, 185, window=radio2, tags="text_tag")

        self.Z_zero_btn = ttk.Button(self.canvas, text="Enter R0 & X0",
                                     command=lambda: self.open_Z_zero_dialog(), style="Custom.TButton")
        SC_canvas.create_window(400, 185, window=self.Z_zero_btn)
        self.Z_zero_btn["state"] = "disabled"

        self.assume_Z_widget = ttk.Entry(self.canvas, textvariable=self.assume_Z_space, width=3,
                                         justify="center", font=('Verdana', int(self.f14.cget('size') + self.ref_f)))
        SC_canvas.create_window(385, 150, window=self.assume_Z_widget, tags="text_tag")

        SC_canvas.create_text(425, 150, text="* Z", tags="text_tag",
                              font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'), justify="center")

        self.SC_text = tk.Text(SC_canvas, bg='#C47B7B',
                               fg='black', font=('Verdana', int(self.f12.cget('size') + self.ref_f), 'bold'))
        SC_canvas.create_window(375, 400, window=self.SC_text)

    def on_radio_change(self):
        # Check the value of `selected_option` to enable/disable the entry and button
        if self.Z_zero_option.get() == 1:
            self.assume_Z_widget.config(state="normal")  # Enable entry widget
            self.Z_zero_btn.config(state="disabled")  # Disable button
        else:
            self.assume_Z_widget.config(state="disabled")  # Disable entry widget
            self.Z_zero_btn.config(state="normal")  # Enable button

    def SC_calc_command(self, SC_type):
        self.SC_type = SC_type
        try:
            self.fault_bus = self.fault_bus_space.get()
            if 0 < self.fault_bus <= self.num_buses:
                Z_zero_radio_option = self.Z_zero_option.get()
                if Z_zero_radio_option == 1:
                    self.Z_product = float(self.assume_Z_widget.get())
                    if self.Z_product <= 0:
                        self.display_error_message("Assumed Z product must be > 0 ")
                        return  # Prevent further execution if Z_product is invalid
                self.display_error_message("")
                self.SC_calc()
        except ValueError:
            self.display_error_message("Please enter a valid Bus number")
            return
        except tk.TclError:
            self.display_error_message("Please enter a valid Bus number")
            return

    def SC_calc(self):
        # Calculate Z values
        Z_values = {}
        Z_fault = 0.001
        Z_zero_radio_option = self.Z_zero_option.get()
        if Z_zero_radio_option == 1:
            for i in range(1, self.num_buses + 1):
                for j in range(i + 1, self.num_buses + 1):
                    pair = f"{i}_{j}"
                    Z_values[pair] = self.Z_product * (self.R_values[pair] + 1j * self.X_values[pair])
            Y_instances = Y_bus_class(self.num_buses)
        else:
            for i in range(1, self.num_buses + 1):
                for j in range(i + 1, self.num_buses + 1):
                    pair = f"{i}_{j}"
                    Z_values[pair] = self.R_Zero_values[pair] + 1j * self.X_Zero_values[pair]
            Y_instances = Y_bus_class(self.num_buses)

        # Add line admittances to the Y-bus matrix
        for pair, Z in Z_values.items():
            i, j = map(int, pair.split('_'))  # Correctly split the pair to extract i and j
            Y = 1 / Z if Z != 0 else 0
            Y_instances.add_line_admittance(i, j, Y, 0, 0)

        Y_zero = Y_instances.get_Y_bus_matrix()
        Y_pos, Y_neg = np.copy(self.Y_bus_matrix), np.copy(self.Y_bus_matrix)

        try:
            Z_zero = np.linalg.inv(Y_zero)
        except np.linalg.LinAlgError:
            Z_zero = 1 / Y_zero

        try:
            Z_pos = np.linalg.inv(Y_pos)
        except np.linalg.LinAlgError:
            Z_pos = 1 / Y_pos

        try:
            Z_neg = np.linalg.inv(Y_neg)
        except np.linalg.LinAlgError:
            Z_neg = 1 / Y_neg

        self.V_output_array = self.V_mag_output_array * np.exp(1j * self.V_ang_output_array)
        P_net = self.PG_output_array - self.PL_array
        Q_net = self.QG_output_array - self.QL_array
        I_node = (P_net + Q_net * 1j) / np.conj(self.V_output_array)
        Y_fault = None
        I_sc_bus = None

        # Calculate fault line currents I_ij_array
        if self.SC_type == "L-L-L fault":
            Z_Thevenin = Z_pos[self.fault_bus - 1][self.fault_bus - 1]

            I_sc_bus = self.V_output_array[self.fault_bus - 1] / (Z_Thevenin + Z_fault)
            Y_fault = np.copy(Y_pos)
            Y_fault[self.fault_bus - 1][self.fault_bus - 1] += 1 / Z_fault

        elif self.SC_type == "L-L fault":
            Z_Thevenin = (Z_pos[self.fault_bus - 1][self.fault_bus - 1] +
                          Z_neg[self.fault_bus - 1][self.fault_bus - 1])

            I_sc_bus = self.V_output_array[self.fault_bus - 1] / (Z_Thevenin + Z_fault)
            Y_fault = Y_pos + Y_neg
            Y_fault[self.fault_bus - 1][self.fault_bus - 1] += 1 / Z_fault

        elif self.SC_type == "L-L-G fault":
            Z1 = Z_pos[self.fault_bus - 1][self.fault_bus - 1]
            Z2 = Z_neg[self.fault_bus - 1][self.fault_bus - 1]
            Z0 = Z_zero[self.fault_bus - 1][self.fault_bus - 1]
            Z_Thevenin = Z1 + 2 * Z2 + Z0  # Adjusted formula
            print("Z_Thevenin", Z_Thevenin)
            I_sc_bus = self.V_output_array[self.fault_bus - 1] / (Z_Thevenin + Z_fault)
            Y_fault = Y_pos + 2 * Y_neg + Y_zero
            Y_fault[self.fault_bus - 1][self.fault_bus - 1] += 1 / Z_fault

        elif self.SC_type == "L-G fault":
            Z1 = Z_pos[self.fault_bus - 1][self.fault_bus - 1]
            Z0 = Z_zero[self.fault_bus - 1][self.fault_bus - 1]
            Z_Thevenin = Z1 + Z0
            print("Z_Thevenin", Z0)

            I_sc_bus = self.V_output_array[self.fault_bus - 1] / (Z_Thevenin + Z_fault)
            Y_fault = Y_pos + Y_zero
            Y_fault[self.fault_bus - 1][self.fault_bus - 1] += 1 / Z_fault

        self.I_sc_bus = I_sc_bus
        # Calculate fault voltages using the inverse of the fault admittance matrix and the source current vector
        try:
            V_fault = np.linalg.inv(Y_fault) @ I_node
        except np.linalg.LinAlgError:
            V_fault = (1 / Y_fault) @ I_node

        # Calculate the line currents under fault conditions using the newly calculated fault voltages
        self.I_ij_fault_array = Y_fault @ (V_fault[:, None] - V_fault)  # Matrix of voltage differences * Y_fault

        # Calculate fault line power flows S_ij_fault_array
        self.S_ij_fault_array = Calc_class.line_power_flows(self.num_buses, V_fault, self.I_ij_fault_array)
        self.SC_text_show()

    def SC_text_show(self):
        self.SC_text["state"] = "normal"
        self.SC_text.delete("1.0", tk.END)
        self.SC_text.insert(tk.END, f"{self.SC_type} at bus #{self.fault_bus}\n")
        I_ij_fault_array_mag = np.abs(self.I_ij_fault_array)
        I_ij_fault_array_ang = np.angle(self.I_ij_fault_array, deg=True)

        self.SC_text.insert(tk.END, f"\nFault Current Isc {np.round(np.abs(self.I_sc_bus), 3)} ∠"
                                    f"{np.round(np.angle(self.I_sc_bus, deg=True), 3)} [p.u.]\n\n")
        self.SC_text.insert(tk.END, "\nLine\tLine Fault Currents [p.u.]\n\n")
        for i in range(self.num_buses):
            for j in range(self.num_buses):
                if i != j and I_ij_fault_array_mag[i][j] != 0 and I_ij_fault_array_ang[i][j] != 0:  # Check if not zero
                    # Access each element with I_ij_array[i][j]
                    element_1 = np.round(I_ij_fault_array_mag[i][j], 3)
                    element_2 = np.round(I_ij_fault_array_ang[i][j], 3)

                    # insert the values in the text widget
                    self.SC_text.insert(tk.END, f"{i + 1}-{j + 1}:\t{element_1} ∠{element_2}°\n")

        S_ij_fault_array_mag = np.abs(self.S_ij_fault_array)
        S_ij_fault_array_ang = np.angle(self.S_ij_fault_array, deg=True)

        self.SC_text.insert(tk.END, "\nLine\tLine Fault Power Flows [p.u.]\n\n")
        for i in range(self.num_buses):
            for j in range(self.num_buses):
                if i != j and S_ij_fault_array_mag[i][j] != 0 and S_ij_fault_array_ang[i][j] != 0:  # Check if not zero
                    # Access each element with I_ij_array[i][j]
                    element_1 = np.round(S_ij_fault_array_mag[i][j], 3)
                    element_2 = np.round(S_ij_fault_array_ang[i][j], 3)

                    # insert the values in the text widget
                    self.SC_text.insert(tk.END, f"{i + 1}-{j + 1}:\t{element_1} ∠{element_2}°\n")

        self.SC_text["state"] = "disabled"

    def open_Z_zero_dialog(self):
        dialog = Z_zero_dialog(self.canvas, "Line Z zero (R & X)")
        self.R_Zero_values = dialog.R_Zero_values
        self.X_Zero_values = dialog.X_Zero_values

    @staticmethod
    def plot_load_gen():
        load = Plots.Load_gen_class()
        load.plot_Load_gen()
        load.plt_show()

    @staticmethod
    def plot_PF():
        PF = Plots.PF_class()
        PF.PF_plot()
        PF.plt_show()

    @staticmethod
    def plot_Voltage_profile():
        V_profile = Plots.V_profile_class()
        V_profile.plot_V_profile()
        V_profile.plt_show()

    @staticmethod
    def plot_Network():
        Network = Plots.PowerNetwork(True)
        Network.visualize_network()

    def PLF_window(self):
        PLF_dialog(self.canvas, "PLF window")

    def pre_save_doc(self):
        self.Save_switch = 0
        self.Round_value = 3
        Data_Store.central_data_store.update_data('Save_switch', self.Save_switch)
        Export_dialog(self.canvas, "Export Results")
        self.Save_switch = Data_Store.central_data_store.get_data('Save_switch')
        self.Round_value = Data_Store.central_data_store.get_data('Round_value')
        if self.Save_switch == 1:
            self.doc_plots()
            threading.Thread(target=self.save_doc, daemon=True).start()

    @staticmethod
    def doc_plots():
        Network = Plots.PowerNetwork(True)
        Network.visualize_network()

        load = Plots.Load_gen_class()
        load.plot_Load_gen()
        load.plt_save()

        PF = Plots.PF_class()
        PF.PF_plot()
        PF.plt_save()

        V_profile = Plots.V_profile_class()
        V_profile.plot_V_profile()
        V_profile.plt_save()

    def save_doc(self):
        current_date = datetime.date.today()
        phrase_text = f"{self.method} Result for {self.num_buses} Buses - {current_date}"
        intro_paragraph = self.doc.add_paragraph()
        intro_run = intro_paragraph.add_run(phrase_text)
        intro_run.bold = True
        intro_run.font.size = Pt(14)  # Set the size as you prefer
        intro_run.font.color.rgb = RGBColor(49, 132, 155)  # Set to blue
        intro_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_heading('Power Flow Overview', level=1)

        self.add_line_parameter_paragraph(f"Iteration Result: {self.iteration_result}")
        self.add_line_parameter_paragraph(f"Total P losses: {np.round(self.total_P_loss, self.Round_value)}")
        self.add_line_parameter_paragraph(f"Total Q losses: {np.round(self.total_Q_loss, self.Round_value)}")
        self.add_line_parameter_paragraph(f"Total P generation:"
                                          f"{np.round(self.total_P_generation, self.Round_value)}")
        self.add_line_parameter_paragraph(f"Total P loads & losses:"
                                          f"{np.round(self.total_P_load_loss, self.Round_value)}")

        # Add a document title as a heading
        self.doc.add_heading('Bus Parameters', level=1)

        # Now add your table under this heading
        headers = ["V", "Ang", "P_G", "Q_G", "P_L", "Q_L", "Bus Type"]
        table = self.doc.add_table(1 + self.num_buses, len(headers))
        table.style = 'Colorful Grid Accent 5'

        # Center text in all cells and populate header row
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Dynamically populate data rows based on num_buses
        for i in range(self.num_buses):
            row_cells = table.row_cells(i + 1)
            row_cells[0].text = f"{np.round(self.V_mag_output_array[i], self.Round_value)}"
            row_cells[1].text = f"{np.round(self.V_ang_output_array[i], self.Round_value)}°"
            row_cells[2].text = f"{np.round(self.PG_output_array[i], self.Round_value)}"
            row_cells[3].text = f"{np.round(self.QG_output_array[i], self.Round_value)}"
            row_cells[4].text = f"{np.round(self.PL_array[i], self.Round_value)}"
            row_cells[5].text = f"{np.round(self.QL_array[i], self.Round_value)}"
            row_cells[6].text = f"{self.bus_type_array[i]}"

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a new heading for Line Parameters
        self.doc.add_heading('Line Parameters', level=1)

        I_ij_array_mag = np.abs(self.I_ij_array)
        I_ij_array_ang = np.angle(self.I_ij_array, deg=True)
        S_ij_array_mag = np.abs(self.S_ij_array)
        S_ij_array_ang = np.angle(self.S_ij_array, deg=True)

        self.add_line_parameter_paragraph("\nLine\tLine Currents [p.u.]")
        for i in range(self.num_buses):
            for j in range(self.num_buses):
                if I_ij_array_mag[i][j] != 0 and I_ij_array_ang[i][j] != 0:
                    element_1 = np.round(I_ij_array_mag[i][j], self.Round_value)
                    element_2 = np.round(I_ij_array_ang[i][j], self.Round_value)
                    content = f"{i + 1}-{j + 1}:\t{element_1} ∠{element_2}°"
                    self.add_line_parameter_paragraph(content)

        self.add_line_parameter_paragraph("\nLine\tLine Power Flows [p.u.]")
        for i in range(self.num_buses):
            for j in range(self.num_buses):
                if S_ij_array_mag[i][j] != 0 and S_ij_array_ang[i][j] != 0:
                    element_1 = np.round(S_ij_array_mag[i][j], self.Round_value)
                    element_2 = np.round(S_ij_array_ang[i][j], self.Round_value)
                    content = f"{i + 1}-{j + 1}:\t{element_1} ∠{element_2}°"
                    self.add_line_parameter_paragraph(content)

        # Adding the heading for the Plots section
        self.doc.add_heading('Plots', level=1)

        # Add the heading and image for each plot
        # 1. Load & Generation Plot
        self.doc.add_heading('Load & Generation Plot', level=2)
        self.doc.add_picture(os.path.join('config', 'Load_gen_plot.png'), width=Inches(6))

        # 2. Network Plot
        self.doc.add_heading('Network Plot', level=2)
        self.doc.add_picture(os.path.join('config', 'network_plot.png'), width=Inches(6))

        # 3. Power Factor Plot
        self.doc.add_heading('Power Factor Plot', level=2)
        self.doc.add_picture(os.path.join('config', 'PF_plot.png'), width=Inches(6))

        # 4. Voltage Profile Plot
        self.doc.add_heading('Voltage Profile Plot', level=2)
        self.doc.add_picture(os.path.join('config', 'V_profile_plot.png'), width=Inches(7))

        self.selected_directory = str(Data_Store.central_data_store.get_data('selected_directory'))
        if not os.path.exists(self.selected_directory):
            os.makedirs(self.selected_directory)

        base_filename = f'{self.method} - {self.num_buses} Buses ({current_date})'
        extension = ".docx"
        filename = base_filename + extension
        full_path = os.path.join(self.selected_directory, filename)
        # Check if the file exists and find a new file name if necessary
        counter = 1
        while os.path.exists(full_path):
            filename = f"{base_filename}({counter}){extension}"
            full_path = os.path.join(self.selected_directory, filename)
            counter += 1

        # Save the document
        self.doc.save(full_path)

        # Open folder
        Results_folder = os.path.join(os.getcwd(), self.selected_directory)
        subprocess.Popen(['explorer', Results_folder])
        self.display_error_message(f"File saved\n[{filename}]")
        self.canvas.itemconfig("error_text", fill="black")

    def add_line_parameter_paragraph(self, content):
        para = self.doc.add_paragraph()
        run = para.add_run(content)
        run.font.size = Pt(12)  # Set the size as you prefer

        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        return para

    def display_error_message(self, message):
        self.canvas.delete("error_text")
        Dark_red = '#8b0000'
        self.canvas.create_text(5 * self.ref_x, 0.5 * self.ref_y, tags="error_text", text=message, fill=Dark_red,
                                font=('Verdana', int(self.f14.cget('size') + self.ref_f), 'bold'))


class Export_dialog(simpledialog.Dialog):
    def __init__(self, parent, title):
        self.Rounding_scale_var = tk.IntVar(value=3)
        self.Save_switch = 0
        self.directory_entry, self.selected_directory_space = None, None
        super().__init__(parent, title)

    def body(self, master):
        self.configure()

        # configure the grid layout
        master.columnconfigure(0, pad=3)
        master.columnconfigure(1, pad=3)

        # Directory field
        directory_label = tk.Label(master)
        directory_label.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)

        self.directory_entry = tk.Entry(master, width=50, bg="white")
        self.directory_entry.grid(row=0, column=1, sticky="we", padx=5, pady=5)
        self.directory_entry.insert(0, "Results")  # Set the default or previously selected directory

        browse_button = tk.Button(master, text="Browse", command=self.browse_directory)
        browse_button.grid(row=0, column=2, padx=5, pady=5)

        # Rounding scale field
        Rounding_scale = tk.Scale(master, from_=0, to=10, resolution=1, orient=tk.HORIZONTAL,
                                  variable=self.Rounding_scale_var, label="Rounding")
        Rounding_scale.grid(row=1, column=1, sticky=tk.EW, padx=5, pady=5)

    def browse_directory(self):
        self.selected_directory_space = filedialog.askdirectory()
        if self.selected_directory_space:  # If a directory was selected
            self.directory_entry.delete(0, tk.END)  # Clear the entry widget
            self.directory_entry.insert(0, self.selected_directory_space)  # Insert into the entry widget

    def apply(self):
        self.Save_switch = 1
        selected_directory = self.directory_entry.get()
        Round_value = self.Rounding_scale_var.get()  # Retrieve the rounding value
        Data_Store.central_data_store.update_data('Round_value', Round_value)
        Data_Store.central_data_store.update_data('selected_directory', selected_directory)
        Data_Store.central_data_store.update_data('Save_switch', self.Save_switch)

    def buttonbox(self):
        box = tk.Frame(self)

        tk.Button(box, text="OK", width=10, command=self.ok, default=tk.ACTIVE).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Button(box, text="Cancel", width=10, command=self.cancel).pack(side=tk.LEFT, padx=5, pady=5)

        self.bind("<Return>", self.ok)
        self.bind("<Escape>", self.cancel)
        box.pack()

    def cancel(self, event=None):
        super().cancel()


class PLF_dialog(simpledialog.Dialog):
    def __init__(self, parent, title):
        self.fonts = Data_Store.central_data_store.get_data('fonts')
        self.num_buses = Data_Store.central_data_store.get_data('num_buses')
        self.num_simulations = Data_Store.central_data_store.get_data('num_simulations')

        self.f12 = self.fonts[1]
        self.bus_space = tk.IntVar()

        super().__init__(parent, title)

    def body(self, master):
        text_frame = tk.Frame(master)
        text_frame.pack()
        ref_f = Data_Store.central_data_store.get_data('ref_f')
        PLF_text = tk.Text(text_frame, bg='#8CA6B8', fg='black', width=50, height=10,
                           font=('Verdana', int(self.f12.cget('size') + ref_f), 'bold'))
        PLF_text.pack(pady=5)

        all_results = Data_Store.central_data_store.get_data('all_results')
        non_convergence_log = Data_Store.central_data_store.get_data('non_convergence_log')
        PLF_text.insert(tk.END, f"Number of Simulations: {self.num_simulations}\n")

        for result in all_results:
            average_magnitude = np.round(result["Average Magnitude"], 3)
            average_angle = np.round(result["Average Angle"], 3)
            lowest_magnitude = np.round(result['Lowest Magnitude'], 3)
            lowest_angle = np.round(result['Lowest Angle'], 3)
            maximum_magnitude = np.round(result['Maximum Magnitude'], 3)
            maximum_angle = np.round(result['Maximum Angle'], 3)
            PLF_text.insert(tk.END, f"Average Magnitude: {average_magnitude}\n")
            PLF_text.insert(tk.END, f"Average Angle: {average_angle}\n")
            PLF_text.insert(tk.END, f"Lowest Magnitude: {lowest_magnitude}\n")
            PLF_text.insert(tk.END, f"Lowest Angle: {lowest_angle}\n")
            PLF_text.insert(tk.END, f"Maximum Magnitude: {maximum_magnitude}\n")
            PLF_text.insert(tk.END, f"Maximum Angle: {maximum_angle}\n")
            PLF_text.insert(tk.END, f"non_convergence_log: {non_convergence_log}\n")
        PLF_text.insert(tk.END, f"non_convergence_%: "
                                f"{(len(non_convergence_log) / self.num_simulations) * 100}%\n")
        PLF_text["state"] = "disabled"

    def buttonbox(self):
        Left_box = tk.Frame(self)
        Left_box.pack(side=tk.LEFT, padx=5, pady=5)
        tk.Button(Left_box, text="Histogram", width=20, command=self.plot_Histogram).pack(side=tk.TOP, padx=5, pady=5)
        tk.Button(Left_box, text="Load/V Correlation", width=20,
                  command=lambda: self.Load_Correlation(None)).pack(side=tk.TOP, padx=5, pady=5)
        tk.Button(Left_box, text="Load/V bus Corr.", width=20,
                  command=lambda: self.Load_Correlation("bus")).pack(side=tk.TOP, padx=5, pady=5)

        middle_box = tk.Frame(self)
        middle_box.pack(side=tk.LEFT, padx=5, pady=5)
        (tk.Button(middle_box, text="Bus Histogram", width=20, command=self.plot_bus_Histogram).
         pack(side=tk.TOP, padx=5, pady=5))
        tk.Button(middle_box, text="Gen/V Correlation", width=20,
                  command=lambda: self.Gen_Correlation(None)).pack(side=tk.TOP, padx=5, pady=5)
        tk.Button(middle_box, text="Gen/V bus Corr.", width=20,
                  command=lambda: self.Gen_Correlation("bus")).pack(side=tk.TOP, padx=5, pady=5)

        Right_box = tk.Frame(self)
        Right_box.pack(side=tk.LEFT, padx=5, pady=5)

        bus_widget = ttk.Entry(Right_box, textvariable=self.bus_space, width=20)
        bus_widget.pack(side=tk.TOP, padx=5, pady=5)
        tk.Button(Right_box, text="3D Correlation", width=20,
                  command=lambda: self.Load_Gen_Correlation(None)).pack(side=tk.TOP, padx=5, pady=5)
        tk.Button(Right_box, text="3D bus Corr.", width=20,
                  command=lambda: self.Load_Gen_Correlation("bus")).pack(side=tk.TOP, padx=5, pady=5)

    @staticmethod
    def plot_Histogram():
        Histogram = Plots.Histogram_class()
        Histogram.Histogram_plot(None)

    def plot_bus_Histogram(self):
        bus_num = int(self.bus_space.get())
        if 0 < bus_num <= self.num_buses:
            Histogram = Plots.Histogram_class()
            Histogram.Histogram_plot(bus_num - 1)

    def Load_Correlation(self, switch):
        if switch:
            bus_num = int(self.bus_space.get())
            if 0 < bus_num <= self.num_buses:
                Correlation = Plots.Correlation_class(bus_num - 1)
                Correlation.Load_Variation(bus_num - 1)
        else:
            Correlation = Plots.Correlation_class(None)
            Correlation.Load_Variation(None)

    def Gen_Correlation(self, switch):
        if switch:
            bus_num = int(self.bus_space.get())
            if 0 < bus_num <= self.num_buses:
                Correlation = Plots.Correlation_class(bus_num - 1)
                Correlation.Generation_Variation(bus_num - 1)
        else:
            Correlation = Plots.Correlation_class(None)
            Correlation.Generation_Variation(None)

    def Load_Gen_Correlation(self, switch):
        if switch:
            bus_num = int(self.bus_space.get())
            if 0 < bus_num <= self.num_buses:
                Correlation = Plots.Correlation_class(bus_num - 1)
                Correlation.Load_Gen_Variation(bus_num - 1)
        else:
            Correlation = Plots.Correlation_class(None)
            Correlation.Load_Gen_Variation(None)

    def cancel(self, event=None):
        super().cancel()


class Z_zero_dialog(simpledialog.Dialog):
    def __init__(self, parent, title):
        self.entries, self.validation_labels, self.scrollable_frame = None, None, None
        self.R_Zero_values, self.X_Zero_values, self.ok_btn = None, None, None
        self.num_buses = Data_Store.central_data_store.get_data('num_buses')
        self.Y_bus_matrix = Data_Store.central_data_store.get_data('Y_bus_matrix')

        fonts = Data_Store.central_data_store.get_data('fonts')
        self.f12 = fonts[1]
        super().__init__(parent, title)

    def body(self, master):
        # Create a canvas and scrollbar within the dialog
        canvas = tk.Canvas(master)
        scrollbar = ttk.Scrollbar(master, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)

        # Configure the canvas
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        self.scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        ref_f = Data_Store.central_data_store.get_data('ref_f')

        # Pack the canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        self.entries = {}
        self.validation_labels = {}
        prefixes = ["R", "X"]
        row = 0

        # Create dynamic entries for the number of buses
        for i in range(1, self.num_buses):
            for j in range(i + 1, self.num_buses + 1):
                if self.Y_bus_matrix[i - 1][j - 1] != 0:
                    tk.Label(self.scrollable_frame, text=f"Line {i}-{j}:",
                             font=('Verdana', int(self.f12.cget('size') + ref_f), 'bold')
                             ).grid(row=row, column=0, sticky=tk.W)
                    col = 1
                    for prefix in prefixes:
                        var_name = f"{prefix}{i}_{j}_space"
                        var = tk.DoubleVar(value=0.0)
                        self.entries[var_name] = var
                        (ttk.Label(self.scrollable_frame, text=f"{prefix}:",
                                   font=('Verdana', int(self.f12.cget('size') + ref_f))).
                         grid(row=row, column=col, sticky=tk.W))
                        col += 1
                        entry = ttk.Entry(self.scrollable_frame, textvariable=var, width=8, justify="center")
                        entry.grid(row=row, column=col, padx=5)
                        col += 1
                        # Adding a label for validation result
                        label = ttk.Label(self.scrollable_frame, text="",
                                          font=('Verdana', int(self.f12.cget('size') + ref_f), 'bold'))
                        label.grid(row=row, column=col, padx=5)
                        self.validation_labels[var_name] = label
                        col += 1
                    row += 1
        return canvas  # Return the canvas as the body of the dialog

    def validate_entry(self):
        # Validation logic here
        all_valid = True  # Start with the assumption that all entries are valid
        for key, var in self.entries.items():
            if self.check_input_error(key):
                self.validation_labels[key].config(text="")
            else:
                self.validation_labels[key].config(text="x", foreground='red')
                all_valid = False

        self.ok_btn["state"] = "normal" if all_valid else "disabled"

    def check_input_error(self, var_name):
        try:
            # Directly attempt to convert the tkinter variable's value to float
            float(self.entries[var_name].get())
            return True
        except ValueError:
            # Catch ValueError if float conversion fails due to incorrect string formats (non-numeric)
            return False
        except tk.TclError:
            # Catch TclError if tkinter variable handling fails (e.g., bad string passed to DoubleVar)
            return False

    def buttonbox(self):
        box = tk.Frame(self)
        tk.Button(box, text="Validate", width=10, command=self.validate_entry).pack(side=tk.LEFT, padx=5, pady=5)
        self.ok_btn = tk.Button(box, text="OK", width=10, command=self.ok)
        self.ok_btn.pack(side=tk.LEFT, padx=5, pady=5)
        self.ok_btn["state"] = "disabled"
        tk.Button(box, text="Cancel", width=10, command=self.cancel).pack(side=tk.LEFT, padx=5, pady=5)
        box.pack()

    def apply(self):
        # Initialize separate dictionaries for R and X values
        self.R_Zero_values = {}
        self.X_Zero_values = {}

        # Process the entry values here, dividing them into separate categories
        for key, var in self.entries.items():
            value = var.get()
            pair = key[1:]  # Extract the pair part, e.g., '1_2' from 'R1_2'
            if key.startswith("R"):
                self.R_Zero_values[pair] = value
            elif key.startswith("X"):
                self.X_Zero_values[pair] = value

    def cancel(self, event=None):
        super().cancel()


class Gen_dialog(simpledialog.Dialog):
    def __init__(self, parent, bus_num, last_P, last_Q, title):
        self.ok_btn = None
        self.total_PG_sub = None
        self.total_QG_sub = None
        self.bus_num = bus_num
        self.last_P = last_P
        self.last_Q = last_Q
        self.entries = {}
        self.validation_labels = {}
        self.init_PG_sub_list = Data_Store.central_data_store.get_data(f"PG_sub_{bus_num}")
        self.init_QG_sub_list = Data_Store.central_data_store.get_data(f"QG_sub_{bus_num}")
        fonts = Data_Store.central_data_store.get_data("fonts")
        self.f12 = fonts[1]
        super().__init__(parent, title=title)

    def body(self, master):
        self.frame = tk.Frame(master)
        self.frame.pack(fill="both", expand=True)
        if self.init_PG_sub_list and self.init_QG_sub_list:
            # Iterate through each index and value from PG and QG lists
            for idx, (pg, qg) in enumerate(zip(self.init_PG_sub_list, self.init_QG_sub_list)):
                self.add_generator_entries(pg, qg, idx)
        else:
            self.add_generator_entries(self.last_P, self.last_Q, 0)
        return self.frame

    def add_generator_entries(self, pg, qg, row):
        pg_var = tk.DoubleVar(value=pg)
        qg_var = tk.DoubleVar(value=qg)

        ttk.Entry(self.frame, textvariable=pg_var, width=8).grid(row=row, column=0, padx=10, pady=5)
        ttk.Entry(self.frame, textvariable=qg_var, width=8).grid(row=row, column=1, padx=10, pady=5)

        # Create validation labels for each entry
        pg_validation_label = tk.Label(self.frame, text="")
        qg_validation_label = tk.Label(self.frame, text="")
        pg_validation_label.grid(row=row, column=3, padx=5)
        qg_validation_label.grid(row=row, column=4, padx=5)

        # Create unique keys for PG and QG values based on the row
        pg_key = f"PG_{self.bus_num}_{row + 1}"
        qg_key = f"QG_{self.bus_num}_{row + 1}"
        self.entries[pg_key] = pg_var
        self.entries[qg_key] = qg_var
        self.validation_labels[pg_key] = pg_validation_label
        self.validation_labels[qg_key] = qg_validation_label

        # Add "+" button only at the first row
        if row == 0:
            add_button = tk.Button(self.frame, text="+",
                                   command=lambda: self.add_generator_entries(0.0, 0.0, len(self.entries) // 2))
            add_button.grid(row=row, column=2, padx=10, pady=5)

    def buttonbox(self):
        box = tk.Frame(self)
        tk.Button(box, text="Validate", width=10, command=self.validate_entry).pack(side=tk.LEFT, padx=5, pady=5)
        self.ok_btn = tk.Button(box, text="OK", width=10, command=self.ok)
        self.ok_btn.pack(side=tk.LEFT, padx=5, pady=5)
        self.ok_btn["state"] = "disabled"
        tk.Button(box, text="Cancel", width=10, command=self.cancel).pack(side=tk.LEFT, padx=5, pady=5)
        box.pack()

    def validate_entry(self):
        all_valid = True
        for key, var in self.entries.items():
            if not self.check_input_error(var):
                self.validation_labels[key].config(text="x", fg='red')
                all_valid = False
            else:
                self.validation_labels[key].config(text="")

        self.ok_btn["state"] = "normal" if all_valid else "disabled"

    @staticmethod
    def check_input_error(value):
        try:
            float(value.get())
            return True
        except ValueError:
            return False
        except tk.TclError:
            return False

    def apply(self):
        PG_sub_list = []
        QG_sub_list = []

        for key, var in self.entries.items():
            if 'PG' in key:
                PG_sub_list.append(var.get())
            elif 'QG' in key:
                QG_sub_list.append(var.get())
        self.total_PG_sub = np.round(sum(PG_sub_list), 4)
        self.total_QG_sub = np.round(sum(QG_sub_list), 4)
        Data_Store.central_data_store.update_data(f"PG_sub_{self.bus_num}", PG_sub_list)
        Data_Store.central_data_store.update_data(f"QG_sub_{self.bus_num}", QG_sub_list)

    def cancel(self, event=None):
        super().cancel()
